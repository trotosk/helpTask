import streamlit as st
import requests
import os
import zipfile
import tempfile
import base64
from pathlib import Path
import json
from datetime import datetime
import numpy as np
from sentence_transformers import SentenceTransformer
import docx
from io import BytesIO
import PyPDF2
import re

# ==================================================
# USUARIOS FIJOS
# ==================================================
USERS = {
    "antonio.alcaraz@softtek.com": "123456",
    "tester@softtek.com": "123456"
}

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.user_email = ""

# PANTALLA DE LOGIN
if not st.session_state.logged_in:
    st.title("🔒 Iniciar sesión")
    email = st.text_input("Correo")
    password = st.text_input("Contraseña", type="password")
    if st.button("Entrar"):
        if email in USERS and USERS[email] == password:
            st.session_state.logged_in = True
            st.session_state.user_email = email
            st.success(f"Bienvenido {email}!")
            st.rerun()
        else:
            st.error("Correo o contraseña incorrectos")
    st.stop()

# ==================================================
# IMPORTS DE TEMPLATES Y CONFIG
# ==================================================
from templates import (
    get_general_template,
    get_code_template,
    get_criterios_Aceptacion_template,
    get_criterios_epica_template,
    get_criterios_mejora_template,
    get_spike_template,
    get_historia_epica_template,
    get_resumen_reunion_template,
    get_criterios_epica_only_history_template,
    get_crear_workitem_template
)

st.set_page_config(
    page_title="Softtek Prompts IA",
    page_icon="🧠",
    layout="wide"
)

API_URL = os.getenv("IA_URL")
API_RESOURCE = os.getenv("IA_RESOURCE_CONSULTA")
TOKEN = os.getenv("IA_TOKEN")

# ==================================================
# ESTADO INICIAL
# ==================================================
defaults = {
    "messages": [],
    "devops_messages": [],
    "doc_messages": [],
    "wiki_messages": [],
    "memory_summary": "",
    # Estado para DevOps
    "devops_incidencias": [],
    "devops_embeddings": None,
    "devops_indexed": False,
    "embedding_model": None,
    "devops_org": "",
    "devops_project": "",
    "devops_pat": "",
    "devops_top_k": 5,
    # Estado para Documentos
    "doc_content": "",
    "doc_chunks": [],
    "doc_embeddings": None,
    "doc_indexed": False,
    "doc_filename": "",
    "doc_top_k": 3,
    "temp_attachments": [],
    "selected_attachment_url": "",
    "selected_attachment_name": "",
    # Estado para Wiki
    "wiki_paginas_contenido": [],
    "wiki_embeddings": None,
    "wiki_referencias": [],
    "wiki_chunks": [],
    "wiki_indexed": False,
    "wiki_top_k": 5,
    "selected_wiki_id": "",
    "selected_wiki_name": "",
    # Estado para Crear Wiki desde Documento
    "wiki_create_doc_content": "",
    "wiki_create_doc_filename": "",
    "wiki_create_estructura_propuesta": None,
    "wiki_create_estructura_editada": None,
    "wiki_create_modo": "nueva",  # "nueva" o "extender"
    "wiki_create_pagina_padre": "",
    "wiki_create_ready_to_create": False
}

for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

CHUNK_SIZE = 10000  # caracteres por fragmento para archivos grandes

# ==================================================
# MAPEO DE CAMPOS POR TIPO DE WORK ITEM
# ==================================================
WORKITEM_FIELD_MAPPING = {
    "User Story": {
        "titulo": {"azure_field": "System.Title", "enabled": True},
        "descripcion": {"azure_field": "System.Description", "enabled": True},
        "acceptance_criteria": {"azure_field": "Microsoft.VSTS.Common.AcceptanceCriteria", "enabled": True},
        "dependencies": {"azure_field": "Custom.Dependencies", "enabled": True},
        "riesgos": {"azure_field": "Custom.Riesgos_US", "enabled": True},
        "value_area": {"azure_field": "Microsoft.VSTS.Common.ValueArea", "enabled": True},
        "team": {"azure_field": "Custom.User_Story_Team", "enabled": True},
        "source": {"azure_field": "Custom.Source", "enabled": True}
    },
    "Feature": {
        "titulo": {"azure_field": "System.Title", "enabled": True},
        "descripcion": {"azure_field": "System.Description", "enabled": True},
        "acceptance_criteria": {"azure_field": "", "enabled": False},  # No existe en Azure
        "dependencies": {"azure_field": "Custom.Dependencias_feature", "enabled": True},
        "riesgos": {"azure_field": "Custom.Riesgos_feature", "enabled": True},
        "value_area": {"azure_field": "Microsoft.VSTS.Common.ValueArea", "enabled": True},
        "team": {"azure_field": "Custom.Team", "enabled": True},
        "source": {"azure_field": "Custom.Source", "enabled": True}
    },
    "Epic": {
        "titulo": {"azure_field": "System.Title", "enabled": True},
        "descripcion": {"azure_field": "System.Description", "enabled": True},
        "acceptance_criteria": {"azure_field": "", "enabled": False},  # No existe en Azure
        "dependencies": {"azure_field": "Custom.Dependencias", "enabled": True},
        "riesgos": {"azure_field": "Custom.Riesgos", "enabled": True},
        "value_area": {"azure_field": "", "enabled": False},  # No existe en Azure
        "team": {"azure_field": "", "enabled": False},  # No existe en Azure
        "source": {"azure_field": "", "enabled": False}  # No existe en Azure
    },
    "Bug": {
        "titulo": {"azure_field": "System.Title", "enabled": True},
        "descripcion": {"azure_field": "System.Description", "enabled": True},
        "acceptance_criteria": {"azure_field": "Microsoft.VSTS.Common.AcceptanceCriteria", "enabled": True},
        "dependencies": {"azure_field": "Custom.Dependencies", "enabled": True},
        "riesgos": {"azure_field": "Custom.Riesgos", "enabled": True},
        "value_area": {"azure_field": "Microsoft.VSTS.Common.ValueArea", "enabled": True},
        "team": {"azure_field": "Custom.Team", "enabled": True},
        "source": {"azure_field": "Custom.Source", "enabled": True}
    },
    "Task": {
        "titulo": {"azure_field": "System.Title", "enabled": True},
        "descripcion": {"azure_field": "System.Description", "enabled": True},
        "acceptance_criteria": {"azure_field": "Microsoft.VSTS.Common.AcceptanceCriteria", "enabled": True},
        "dependencies": {"azure_field": "Custom.Dependencies", "enabled": True},
        "riesgos": {"azure_field": "Custom.Riesgos", "enabled": True},
        "value_area": {"azure_field": "Microsoft.VSTS.Common.ValueArea", "enabled": True},
        "team": {"azure_field": "Custom.Team", "enabled": True},
        "source": {"azure_field": "Custom.Source", "enabled": True}
    }
}

# ==================================================
# HELPERS ORIGINALES
# ==================================================
def call_ia(payload):
    headers = {
        "Authorization": f"Bearer {TOKEN}",
        "Content-Type": "application/json"
    }
    r = requests.post(
        API_URL + API_RESOURCE,
        json=payload,
        headers=headers,
        timeout=120
    )
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

def get_template(tipo):
    return {
        "Libre": get_general_template(),
        "PO Casos exito": get_criterios_Aceptacion_template(),
        "Programador Python": get_code_template(),
        "PO Definicion epica": get_criterios_epica_template(),
        "PO Definicion epica una historia": get_criterios_epica_only_history_template(),
        "PO Definicion mejora tecnica": get_criterios_mejora_template(),
        "PO Definicion spike": get_spike_template(),
        "PO Definicion historia": get_historia_epica_template(),
        "PO resumen reunion": get_resumen_reunion_template(),
        "PO Crear Work Item": get_crear_workitem_template()
    }.get(tipo, get_general_template())

def sanitize_json_string(json_str):
    """Fix unescaped double quotes inside JSON string values (common in AI-generated HTML content)."""
    result = []
    in_string = False
    i = 0
    n = len(json_str)
    while i < n:
        c = json_str[i]
        # Escaped character: pass both chars through unchanged
        if c == '\\' and i + 1 < n:
            result.append(c)
            result.append(json_str[i + 1])
            i += 2
            continue
        if c == '"':
            if not in_string:
                in_string = True
                result.append(c)
            else:
                # Look ahead past whitespace to decide if this closes the string
                j = i + 1
                while j < n and json_str[j] in ' \t\n\r':
                    j += 1
                next_c = json_str[j] if j < n else None
                if next_c in (',', ':', '}', ']', None):
                    # Closing quote
                    in_string = False
                    result.append(c)
                else:
                    # Unescaped quote inside string value — escape it
                    result.append('\\')
                    result.append(c)
        else:
            result.append(c)
        i += 1
    return ''.join(result)

def add_table_borders_to_html(html):
    """Añade estilos CSS inline a las tablas HTML para que se vean con bordes en Azure DevOps."""
    if not html:
        return html
    # Solo actúa si hay tablas en el contenido
    if '<table' not in html.lower():
        return html
    import re
    # <table> — solo si aún no tiene style=
    html = re.sub(
        r'<table(?![^>]*style=)',
        '<table style="border-collapse:collapse;width:100%;"',
        html, flags=re.IGNORECASE
    )
    # <th> — cabeceras con fondo gris
    html = re.sub(
        r'<th(?![^>]*style=)',
        '<th style="border:1px solid #666;padding:6px 8px;background-color:#f2f2f2;text-align:left;"',
        html, flags=re.IGNORECASE
    )
    # <td> — celdas normales
    html = re.sub(
        r'<td(?![^>]*style=)',
        '<td style="border:1px solid #666;padding:6px 8px;"',
        html, flags=re.IGNORECASE
    )
    return html

def resumir_conversacion(messages):
    resumen_prompt = [
        {"role": "system", "content": "Resume la conversación técnica manteniendo contexto y decisiones"},
        {"role": "user", "content": "\n".join(f"{m['role']}: {m['content']}" for m in messages)}
    ]
    return call_ia({"model": st.session_state.model, "messages": resumen_prompt})

# ==================================================
# HELPERS PARA AZURE DEVOPS
# ==================================================

@st.cache_resource
def cargar_modelo_embeddings():
    """Carga el modelo de embeddings una sola vez"""
    return SentenceTransformer('all-MiniLM-L6-v2')

def obtener_incidencias_devops(organization, project, pat, area_path=None, work_item_types=None, max_items=400):
    """
    Obtiene work items de Azure DevOps con filtros configurables
    """
    url = f"https://dev.azure.com/{organization}/{project}/_apis/wit/wiql?api-version=7.1"
    
    # Construir filtro de tipos
    if not work_item_types or len(work_item_types) == 0:
        work_item_types = ['Bug']
    
    if len(work_item_types) == 1:
        type_filter = f"[System.WorkItemType] = '{work_item_types[0]}'"
    else:
        types_str = "', '".join(work_item_types)
        type_filter = f"[System.WorkItemType] IN ('{types_str}')"
    
    # Construir filtro de área
    area_filter = ""
    if area_path:
        area_filter = f"AND [System.AreaPath] = '{area_path}'"
    
    wiql = {
        "query": f"""
            SELECT [System.Id], [System.Title], [System.State], 
                   [System.Description], [System.Tags], 
                   [System.WorkItemType], [System.AreaPath],
                   [Microsoft.VSTS.Common.ResolvedReason],
                   [System.CreatedDate], [System.ChangedDate]
            FROM WorkItems 
            WHERE {type_filter}
            AND [System.State] <> 'Removed'
            {area_filter}
            ORDER BY [System.ChangedDate] DESC
        """
    }
    
    # Encoding del PAT
    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Basic {encoded_credentials}"
    }
    
    try:
        # Debug
        st.info(f"🔍 Consultando: {organization}/{project}")
        st.info(f"📋 Tipos: {', '.join(work_item_types)}")
        if area_path:
            st.info(f"📁 Área: {area_path}")
        st.info(f"🔢 Límite: {max_items} items")
        
        response = requests.post(url, json=wiql, headers=headers, timeout=30)
        st.write(f"**Status Code:** {response.status_code}")
        
        if response.status_code != 200:
            st.error(f"❌ Error HTTP {response.status_code}")
            st.code(response.text[:500])
            return []
        
        response.raise_for_status()
        
        try:
            response_json = response.json()
        except json.JSONDecodeError as e:
            st.error(f"❌ Error al parsear JSON: {str(e)}")
            st.code(response.text[:500])
            return []
        
        work_item_ids = [item["id"] for item in response_json.get("workItems", [])]
        
        if not work_item_ids:
            st.warning("⚠️ La query no devolvió ningún Work Item")
            st.info("Verifica que existan items del tipo seleccionado")
            return []
        
        # Limitar al máximo configurado
        work_item_ids = work_item_ids[:max_items]
        st.success(f"✅ Se encontraron {len(work_item_ids)} work items")
        
        # Obtener detalles en lotes de 200
        all_items = []
        batch_size = 200
        
        for i in range(0, len(work_item_ids), batch_size):
            batch_ids = work_item_ids[i:i+batch_size]
            ids_str = ",".join(map(str, batch_ids))
            details_url = f"https://dev.azure.com/{organization}/{project}/_apis/wit/workitems?ids={ids_str}&api-version=7.1"
            
            details_response = requests.get(details_url, headers=headers, timeout=30)
            details_response.raise_for_status()
            
            for item in details_response.json().get("value", []):
                fields = item.get("fields", {})
                all_items.append({
                    "id": item["id"],
                    "tipo": fields.get("System.WorkItemType", ""),
                    "titulo": fields.get("System.Title", "Sin título"),
                    "descripcion": fields.get("System.Description", "Sin descripción"),
                    "estado": fields.get("System.State", ""),
                    "area": fields.get("System.AreaPath", ""),
                    "tags": fields.get("System.Tags", ""),
                    "resolucion": fields.get("Microsoft.VSTS.Common.ResolvedReason", ""),
                    "fecha_creacion": fields.get("System.CreatedDate", ""),
                    "fecha_cambio": fields.get("System.ChangedDate", ""),
                    "url": item.get("url", "")
                })
        
        return all_items
    
    except requests.exceptions.Timeout:
        st.error("❌ Timeout: Azure DevOps no respondió a tiempo")
        return []
    except requests.exceptions.RequestException as e:
        st.error(f"❌ Error de conexión: {str(e)}")
        if hasattr(e.response, 'text'):
            st.code(e.response.text[:500])
        return []

def obtener_attachments_workitem(organization, project, pat, work_item_id):
    """
    Obtiene la lista de attachments de un work item
    """
    url = f"https://dev.azure.com/{organization}/{project}/_apis/wit/workitems/{work_item_id}?$expand=all&api-version=7.1"
    
    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Basic {encoded_credentials}"
    }
    
    try:
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        work_item = response.json()
        relations = work_item.get("relations", [])
        
        attachments = []
        for relation in relations:
            if relation.get("rel") == "AttachedFile":
                file_name = relation.get("attributes", {}).get("name", "Unknown")
                # Solo incluir archivos .docx
                if file_name.lower().endswith('.docx'):
                    attachments.append({
                        "url": relation.get("url"),
                        "name": file_name
                    })
        
        return attachments
    
    except Exception as e:
        st.error(f"Error al obtener attachments: {str(e)}")
        return []

def descargar_attachment_devops(attachment_url, pat):
    """
    Descarga un attachment desde Azure DevOps
    """
    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()
    
    headers = {
        "Authorization": f"Basic {encoded_credentials}"
    }
    
    try:
        response = requests.get(attachment_url, headers=headers, timeout=30)
        response.raise_for_status()
        return response.content
    except Exception as e:
        st.error(f"Error al descargar attachment: {str(e)}")
        return None

def crear_workitem_devops(organization, project, pat, work_item_type, campos, field_mappings):
    """
    Crea un work item en Azure DevOps con nombres de campos personalizables

    Args:
        organization: Organización de Azure DevOps
        project: Proyecto de Azure DevOps
        pat: Personal Access Token
        work_item_type: Tipo de work item (Bug, User Story, Task, Feature, Epic)
        campos: Diccionario con los valores de los campos del work item
        field_mappings: Diccionario con el mapeo de nombres de campos locales a Azure
            Formato: {'campo_local': {'azure_field': 'Azure.Field.Name', 'enabled': True/False, 'value': 'valor'}}

    Returns:
        Diccionario con 'success', 'id', 'url', 'error'
    """
    url = f"https://dev.azure.com/{organization}/{project}/_apis/wit/workitems/${work_item_type}?api-version=7.1"

    # Encoding del PAT
    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()

    headers = {
        "Content-Type": "application/json-patch+json",
        "Authorization": f"Basic {encoded_credentials}"
    }

    # Construir el body en formato JSON Patch
    # Azure DevOps requiere un array de operaciones de tipo "add"
    body = []

    # Campos HTML que deben llevar bordes en las tablas
    HTML_FIELDS = {'descripcion', 'acceptance_criteria', 'dependencies', 'riesgos'}

    # Recorrer todos los campos y agregar los que estén habilitados y tengan valor
    for field_name, field_value in campos.items():
        # Verificar si el campo tiene configuración en field_mappings
        if field_name in field_mappings:
            mapping = field_mappings[field_name]

            # Solo agregar si está habilitado, tiene un campo de Azure definido, y tiene valor
            if mapping.get('enabled', False) and mapping.get('azure_field') and field_value:
                # Inyectar estilos de borde en tablas HTML antes de enviar a Azure DevOps
                if field_name in HTML_FIELDS:
                    field_value = add_table_borders_to_html(field_value)
                body.append({
                    "op": "add",
                    "path": f"/fields/{mapping['azure_field']}",
                    "value": field_value
                })

    # Area e Iteration son casos especiales (siempre System.AreaPath y System.IterationPath)
    if campos.get('area_path'):
        body.append({
            "op": "add",
            "path": "/fields/System.AreaPath",
            "value": campos['area_path']
        })

    if campos.get('iteration_path'):
        body.append({
            "op": "add",
            "path": "/fields/System.IterationPath",
            "value": campos['iteration_path']
        })

    try:
        response = requests.post(url, json=body, headers=headers, timeout=30)

        if response.status_code == 200 or response.status_code == 201:
            data = response.json()
            work_item_id = data.get('id')
            work_item_url = f"https://dev.azure.com/{organization}/{project}/_workitems/edit/{work_item_id}"

            return {
                'success': True,
                'id': work_item_id,
                'url': work_item_url,
                'error': None
            }
        else:
            error_msg = f"Error HTTP {response.status_code}"
            try:
                error_data = response.json()
                if 'message' in error_data:
                    error_msg += f": {error_data['message']}"
            except:
                error_msg += f": {response.text[:200]}"

            return {
                'success': False,
                'id': None,
                'url': None,
                'error': error_msg
            }

    except Exception as e:
        return {
            'success': False,
            'id': None,
            'url': None,
            'error': str(e)
        }

def limpiar_html(texto):
    """Limpia tags HTML básicos del texto"""
    if not texto:
        return ""
    import re
    texto = re.sub(r'<[^>]+>', ' ', texto)
    texto = re.sub(r'\s+', ' ', texto)
    return texto.strip()

def generar_embeddings_incidencias(incidencias, modelo):
    """
    Genera embeddings para cada incidencia
    """
    textos = []
    for inc in incidencias:
        texto_completo = f"{inc['titulo']} {limpiar_html(inc['descripcion'])} {inc['tags']} {inc['resolucion']}"
        textos.append(texto_completo)
    
    with st.spinner("🔄 Generando embeddings de incidencias..."):
        embeddings = modelo.encode(textos, show_progress_bar=True)
    
    return np.array(embeddings)

def buscar_incidencias_similares(query, incidencias, embeddings, modelo, top_k=5):
    """
    Busca las incidencias más similares a la query usando embeddings
    """
    query_embedding = modelo.encode([query])[0]
    
    similitudes = np.dot(embeddings, query_embedding) / (
        np.linalg.norm(embeddings, axis=1) * np.linalg.norm(query_embedding)
    )
    
    top_indices = np.argsort(similitudes)[-top_k:][::-1]
    
    resultados = []
    for idx in top_indices:
        resultados.append({
            "incidencia": incidencias[idx],
            "similitud": float(similitudes[idx])
        })
    
    return resultados

def construir_contexto_devops(incidencias_similares):
    """
    Construye el contexto para enviar a la IA con las incidencias encontradas
    """
    contexto = "**Work items similares encontrados en Azure DevOps:**\n\n"
    
    for i, resultado in enumerate(incidencias_similares, 1):
        inc = resultado["incidencia"]
        sim = resultado["similitud"]
        
        contexto += f"**Work Item #{i}** (Similitud: {sim:.2%})\n"
        contexto += f"- **ID**: {inc['id']}\n"
        contexto += f"- **Tipo**: {inc['tipo']}\n"
        contexto += f"- **Título**: {inc['titulo']}\n"
        contexto += f"- **Estado**: {inc['estado']}\n"
        if inc['resolucion']:
            contexto += f"- **Resolución**: {inc['resolucion']}\n"
        contexto += f"- **Descripción**: {limpiar_html(inc['descripcion'])[:500]}...\n"
        contexto += f"- **Tags**: {inc['tags']}\n\n"
    
    return contexto

# ==================================================
# HELPERS PARA DOCUMENTOS
# ==================================================

def leer_docx_desde_bytes(file_bytes):
    """Lee un documento Word desde bytes"""
    try:
        doc = docx.Document(BytesIO(file_bytes))
        texto_completo = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                texto_completo.append(para.text.strip())
        
        # También extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texto_completo.append(cell.text.strip())
        
        return "\n\n".join(texto_completo)
    except Exception as e:
        st.error(f"Error al leer documento: {str(e)}")
        return ""

def descargar_documento_url(url):
    """Descarga un documento desde una URL pública"""
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        return response.content
    except Exception as e:
        st.error(f"Error al descargar documento: {str(e)}")
        return None

def dividir_en_chunks(texto, chunk_size=1000):
    """Divide el texto en fragmentos para embeddings"""
    # Dividir por párrafos primero
    paragrafos = [p.strip() for p in texto.split('\n\n') if p.strip()]
    
    chunks = []
    chunk_actual = ""
    
    for para in paragrafos:
        if len(chunk_actual) + len(para) < chunk_size:
            chunk_actual += para + "\n\n"
        else:
            if chunk_actual:
                chunks.append(chunk_actual.strip())
            chunk_actual = para + "\n\n"
    
    if chunk_actual:
        chunks.append(chunk_actual.strip())
    
    return chunks

def generar_embeddings_documento(chunks, modelo):
    """Genera embeddings para los chunks del documento"""
    with st.spinner("🔄 Generando embeddings del documento..."):
        embeddings = modelo.encode(chunks, show_progress_bar=True)
    return np.array(embeddings)

def buscar_chunks_similares(query, chunks, embeddings, modelo, top_k=3):
    """Busca los chunks más relevantes del documento"""
    query_embedding = modelo.encode([query])[0]
    
    similitudes = np.dot(embeddings, query_embedding) / (
        np.linalg.norm(embeddings, axis=1) * np.linalg.norm(query_embedding)
    )
    
    top_indices = np.argsort(similitudes)[-top_k:][::-1]
    
    resultados = []
    for idx in top_indices:
        resultados.append({
            "chunk": chunks[idx],
            "similitud": float(similitudes[idx]),
            "indice": idx
        })
    
    return resultados

def construir_contexto_documento(chunks_similares):
    """Construye el contexto para enviar a Frida con los chunks relevantes"""
    contexto = "**Fragmentos relevantes del documento:**\n\n"
    
    for i, resultado in enumerate(chunks_similares, 1):
        sim = resultado["similitud"]
        chunk = resultado["chunk"]
        
        contexto += f"**Fragmento #{i}** (Relevancia: {sim:.2%})\n"
        contexto += f"{chunk}\n\n"
        contexto += "---\n\n"
    
    return contexto

# ==================================================
# HELPERS PARA AZURE DEVOPS WIKI
# ==================================================

def obtener_wikis_proyecto(organization, project, pat):
    """
    Obtiene la lista de wikis disponibles en el proyecto
    """
    url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis?api-version=7.1"

    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Basic {encoded_credentials}"
    }

    try:
        response = requests.get(url, headers=headers, timeout=30)

        if response.status_code == 401:
            st.error("❌ Error 401: No autorizado para acceder a las Wikis")
            st.warning("""
            **El PAT necesita estos permisos:**
            - ✅ **Wiki (Read)** - Para leer wikis
            - O alternativamente: **Code (Read)** - Da acceso a repos y wikis

            **Pasos para verificar/añadir permisos:**
            1. Ve a Azure DevOps → User Settings → Personal Access Tokens
            2. Edita tu PAT o crea uno nuevo
            3. En los scopes, selecciona: **Wiki (Read)** o **Code (Read)**
            4. Guarda y usa el nuevo PAT en la configuración
            """)
            return []

        response.raise_for_status()

        wikis_data = response.json()
        wikis = wikis_data.get("value", [])

        return [{
            "id": wiki.get("id"),
            "name": wiki.get("name"),
            "type": wiki.get("type"),
            "url": wiki.get("url")
        } for wiki in wikis]

    except requests.exceptions.RequestException as e:
        if hasattr(e, 'response') and e.response is not None:
            st.error(f"❌ Error HTTP {e.response.status_code}: {str(e)}")
        else:
            st.error(f"❌ Error de conexión: {str(e)}")
        return []
    except Exception as e:
        st.error(f"❌ Error inesperado: {str(e)}")
        return []

def obtener_subpaginas_especificas(organization, project, pat, wiki_id, page_path):
    """
    Obtiene las subpáginas de una página específica
    """
    import urllib.parse
    path_encoded = urllib.parse.quote(page_path, safe='')
    url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki_id}/pages?path={path_encoded}&recursionLevel=5&includeContent=false&api-version=7.1"

    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Basic {encoded_credentials}"
    }

    try:
        response = requests.get(url, headers=headers, timeout=30)

        if response.status_code != 200:
            return []

        data = response.json()

        # Función para aplanar subpáginas
        def aplanar_subpaginas(page):
            paginas = []

            # Procesar subpáginas si existen
            if "subPages" in page and page["subPages"]:
                for subpage in page["subPages"]:
                    page_path = subpage.get('path', '')
                    if page_path:
                        paginas.append({
                            "id": subpage.get("id", page_path),
                            "path": page_path,
                            "order": subpage.get("order", 0),
                            "gitItemPath": subpage.get("gitItemPath", ""),
                            "url": subpage.get("url", ""),
                            "isParentPage": subpage.get("isParentPage", False)
                        })
                    # Recursivamente obtener subpáginas de subpáginas
                    paginas.extend(aplanar_subpaginas(subpage))

            return paginas

        return aplanar_subpaginas(data)

    except Exception as e:
        st.warning(f"⚠️ Error al obtener subpáginas de {page_path}: {str(e)}")
        return []

def obtener_paginas_wiki(organization, project, pat, wiki_id, recursion_level=5):
    """
    Obtiene la lista de páginas de una wiki con toda su jerarquía
    recursion_level: Niveles de profundidad a obtener (default 5 para obtener toda la estructura)
    """
    url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki_id}/pages?recursionLevel={recursion_level}&api-version=7.1"

    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Basic {encoded_credentials}"
    }

    try:
        response = requests.get(url, headers=headers, timeout=30)

        if response.status_code == 401:
            st.error("❌ Error 401: No autorizado para acceder a las páginas de la Wiki")
            st.info("Verifica que tu PAT tenga permisos de **Wiki (Read)** o **Code (Read)**")
            return []

        response.raise_for_status()

        data = response.json()

        # DEBUG: Mostrar estructura de respuesta
        with st.expander("🔍 DEBUG - Ver respuesta de la API", expanded=False):
            st.write("**Status Code:**", response.status_code)
            st.write("**Estructura de la respuesta:**")
            st.json(data)
            if data:
                st.write("**Claves en el nivel raíz:**", list(data.keys()))

        # Verificar si hay páginas
        if not data:
            st.warning("⚠️ La respuesta de la API está vacía")
            return []

        # Función recursiva para aplanar la estructura de páginas
        def aplanar_paginas(page, nivel=0):
            paginas = []

            # Añadir página actual si tiene path (excluyendo la raíz "/")
            # Usamos path como identificador ya que no todas las respuestas tienen 'id'
            page_path = page.get('path', '')
            if page_path and page_path != '/':
                paginas.append({
                    "id": page.get("id", page_path),  # Usar path si no hay id
                    "path": page_path,
                    "order": page.get("order", 0),
                    "gitItemPath": page.get("gitItemPath", ""),
                    "url": page.get("url", ""),
                    "isParentPage": page.get("isParentPage", False)
                })

            # Procesar subpáginas si existen
            if "subPages" in page and page["subPages"]:
                for subpage in page["subPages"]:
                    paginas.extend(aplanar_paginas(subpage, nivel + 1))

            return paginas

        # Si hay páginas, aplanarlas
        if "path" in data or "subPages" in data:
            paginas_encontradas = aplanar_paginas(data)

            if paginas_encontradas:
                st.success(f"✅ Se encontraron {len(paginas_encontradas)} página(s)")
            else:
                st.warning("⚠️ Se procesó la respuesta pero no se encontraron páginas válidas")

            return paginas_encontradas
        else:
            st.warning("⚠️ La respuesta no tiene 'path' ni 'subPages' en la raíz")
            st.write("**Claves disponibles:**", list(data.keys()) if isinstance(data, dict) else "No es un diccionario")
            return []

    except requests.exceptions.RequestException as e:
        if hasattr(e, 'response') and e.response is not None:
            st.error(f"❌ Error HTTP {e.response.status_code}: {str(e)}")
        else:
            st.error(f"❌ Error de conexión: {str(e)}")
        return []
    except Exception as e:
        st.error(f"❌ Error inesperado al obtener páginas: {str(e)}")
        return []

def obtener_contenido_pagina_wiki(organization, project, pat, wiki_id, page_id):
    """
    Obtiene el contenido de una página específica de la wiki
    page_id puede ser el ID numérico o el path de la página
    """
    # Si page_id es un path (empieza con /), usar parámetro path
    if isinstance(page_id, str) and page_id.startswith('/'):
        url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki_id}/pages?path={page_id}&includeContent=true&api-version=7.1"
    else:
        # Si es un ID numérico, usar la ruta tradicional
        url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki_id}/pages/{page_id}?includeContent=true&api-version=7.1"

    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Basic {encoded_credentials}"
    }

    try:
        response = requests.get(url, headers=headers, timeout=30)

        if response.status_code == 401:
            st.error(f"❌ Error 401: No autorizado para acceder al contenido de la página")
            st.info("Verifica que tu PAT tenga permisos de **Wiki (Read)** o **Code (Read)**")
            return None

        response.raise_for_status()

        data = response.json()

        return {
            "id": data.get("id"),
            "path": data.get("path"),
            "content": data.get("content", ""),
            "gitItemPath": data.get("gitItemPath", "")
        }

    except requests.exceptions.RequestException as e:
        if hasattr(e, 'response') and e.response is not None:
            st.error(f"❌ Error HTTP {e.response.status_code} al obtener contenido")
        else:
            st.error(f"❌ Error de conexión: {str(e)}")
        return None
    except Exception as e:
        st.error(f"❌ Error inesperado al obtener contenido: {str(e)}")
        return None

def limpiar_markdown(texto):
    """
    Limpia tags markdown y deja texto limpio
    """
    if not texto:
        return ""
    import re
    # Eliminar bloques de código
    texto = re.sub(r'```[\s\S]*?```', '', texto)
    # Eliminar código inline
    texto = re.sub(r'`[^`]+`', '', texto)
    # Eliminar imágenes
    texto = re.sub(r'!\[.*?\]\(.*?\)', '', texto)
    # Eliminar links pero mantener texto
    texto = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', texto)
    # Eliminar headers markdown
    texto = re.sub(r'^#+\s+', '', texto, flags=re.MULTILINE)
    # Eliminar énfasis
    texto = re.sub(r'[*_]{1,2}([^*_]+)[*_]{1,2}', r'\1', texto)
    # Limpiar espacios múltiples
    texto = re.sub(r'\s+', ' ', texto)
    return texto.strip()

def generar_embeddings_wiki(paginas_contenido, modelo):
    """
    Genera embeddings para las páginas de la wiki
    paginas_contenido: lista de dict con 'path', 'chunks', etc.
    """
    todos_chunks = []
    referencias = []  # Para mantener referencia de página y chunk

    for pagina in paginas_contenido:
        for idx, chunk in enumerate(pagina['chunks']):
            todos_chunks.append(chunk)
            referencias.append({
                'path': pagina['path'],
                'chunk_idx': idx,
                'page_id': pagina['id']
            })

    with st.spinner("🔄 Generando embeddings de páginas Wiki..."):
        embeddings = modelo.encode(todos_chunks, show_progress_bar=True)

    return np.array(embeddings), referencias

def buscar_chunks_wiki_similares(query, chunks, embeddings, referencias, modelo, top_k=5):
    """
    Busca los chunks más relevantes de las páginas Wiki
    """
    query_embedding = modelo.encode([query])[0]

    similitudes = np.dot(embeddings, query_embedding) / (
        np.linalg.norm(embeddings, axis=1) * np.linalg.norm(query_embedding)
    )

    top_indices = np.argsort(similitudes)[-top_k:][::-1]

    resultados = []
    for idx in top_indices:
        resultados.append({
            "chunk": chunks[idx],
            "similitud": float(similitudes[idx]),
            "path": referencias[idx]['path'],
            "page_id": referencias[idx]['page_id'],
            "chunk_idx": referencias[idx]['chunk_idx']
        })

    return resultados

def construir_contexto_wiki(chunks_similares):
    """
    Construye el contexto para enviar a Frida con los chunks relevantes de la Wiki
    """
    contexto = "**Fragmentos relevantes de la Wiki de Azure DevOps:**\n\n"

    for i, resultado in enumerate(chunks_similares, 1):
        sim = resultado["similitud"]
        chunk = resultado["chunk"]
        path = resultado["path"]

        contexto += f"**Página: {path}** - Fragmento #{i} (Relevancia: {sim:.2%})\n"
        contexto += f"{chunk}\n\n"
        contexto += "---\n\n"

    return contexto

# ==================================================
# HELPERS PARA CREACIÓN DE WIKI DESDE DOCUMENTOS
# ==================================================

def leer_pdf_desde_bytes(file_bytes):
    """Lee un documento PDF desde bytes"""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        texto_completo = []

        for page in pdf_reader.pages:
            texto = page.extract_text()
            if texto.strip():
                texto_completo.append(texto.strip())

        return "\n\n".join(texto_completo)
    except Exception as e:
        st.error(f"Error al leer PDF: {str(e)}")
        return ""

def detectar_encabezados_principales(contenido_documento):
    """
    Detecta automáticamente los encabezados principales del documento
    Retorna lista de encabezados con su posición
    """
    lineas = contenido_documento.split('\n')
    encabezados = []

    for idx, linea in enumerate(lineas):
        linea_limpia = linea.strip()

        # Detectar encabezados por patrones comunes
        es_encabezado = False

        # Patrón 1: Números al inicio (1., 2., 1.1, etc.)
        if re.match(r'^\d+\.(\d+\.)*\s+[A-Z]', linea_limpia):
            es_encabezado = True

        # Patrón 2: Todo en mayúsculas y largo suficiente
        elif len(linea_limpia) > 5 and len(linea_limpia) < 100 and linea_limpia.isupper() and not linea_limpia.startswith('-'):
            es_encabezado = True

        # Patrón 3: Markdown headers (#, ##, ###)
        elif re.match(r'^#{1,3}\s+', linea_limpia):
            es_encabezado = True

        # Patrón 4: Línea seguida de guiones/iguales (estilo rst)
        elif idx < len(lineas) - 1:
            siguiente = lineas[idx + 1].strip()
            if len(linea_limpia) > 3 and len(linea_limpia) < 100:
                if re.match(r'^[=\-]{3,}$', siguiente):
                    es_encabezado = True

        if es_encabezado and linea_limpia:
            # Limpiar el título
            titulo = re.sub(r'^\d+\.(\d+\.)*\s*', '', linea_limpia)
            titulo = re.sub(r'^#+\s*', '', titulo)
            titulo = titulo.strip()

            if titulo and len(titulo) > 3:
                encabezados.append({
                    'titulo': titulo,
                    'idx': idx,
                    'linea_original': linea_limpia
                })

    return encabezados

def dividir_documento_por_encabezados(contenido_documento, filename):
    """
    Divide el documento en secciones por encabezados detectados
    Retorna estructura lista para crear wiki
    """
    encabezados = detectar_encabezados_principales(contenido_documento)

    if not encabezados or len(encabezados) < 2:
        # Si no hay encabezados, devolver documento completo
        return {
            "paginas": [
                {
                    "titulo": filename.replace('.docx', '').replace('.pdf', ''),
                    "es_raiz": True,
                    "padre": None,
                    "contenido_markdown": f"# {filename}\n\n{contenido_documento}",
                    "orden": 0
                }
            ]
        }

    lineas = contenido_documento.split('\n')
    paginas = []

    # Página índice
    doc_titulo = filename.replace('.docx', '').replace('.pdf', '')
    indice_contenido = f"# {doc_titulo}\n\n## Índice\n\n"
    for enc in encabezados[:15]:
        indice_contenido += f"- {enc['titulo']}\n"

    paginas.append({
        "titulo": "Índice",
        "es_raiz": True,
        "padre": None,
        "contenido_markdown": indice_contenido,
        "orden": 0
    })

    # Una página por cada encabezado con contenido literal
    for i, encabezado in enumerate(encabezados):
        inicio = encabezado['idx']
        fin = encabezados[i + 1]['idx'] if i < len(encabezados) - 1 else len(lineas)

        # Contenido literal de la sección
        contenido_seccion = '\n'.join(lineas[inicio:fin])
        contenido_markdown = f"# {encabezado['titulo']}\n\n{contenido_seccion.strip()}"

        paginas.append({
            "titulo": encabezado['titulo'][:50],
            "es_raiz": False,
            "padre": "Índice",
            "contenido_markdown": contenido_markdown,
            "orden": i + 1
        })

    return {"paginas": paginas}

def extraer_contenido_seccion(contenido_documento, seccion_origen, titulo_pagina):
    """
    Extrae el contenido completo de una sección específica del documento
    """
    if not seccion_origen:
        return f"# {titulo_pagina}\n\n[Contenido pendiente de asignar]"

    # Buscar por los encabezados posibles
    posibles_encabezados = [s.strip() for s in seccion_origen.split('|')]

    # Dividir el documento en líneas
    lineas = contenido_documento.split('\n')

    # Buscar el inicio de la sección
    inicio_idx = None
    for idx, linea in enumerate(lineas):
        linea_limpia = linea.strip()
        for encabezado in posibles_encabezados:
            if encabezado.lower() in linea_limpia.lower():
                inicio_idx = idx
                break
        if inicio_idx is not None:
            break

    if inicio_idx is None:
        # Si no se encuentra, devolver una porción del documento
        return f"# {titulo_pagina}\n\n{contenido_documento[:3000]}"

    # Buscar el final de la sección (siguiente encabezado principal o final del documento)
    fin_idx = len(lineas)
    for idx in range(inicio_idx + 1, len(lineas)):
        linea = lineas[idx].strip()
        # Detectar encabezados (números, mayúsculas, etc.)
        if (linea and (
            re.match(r'^\d+\.', linea) or  # Empieza con número
            (len(linea) > 10 and linea.isupper()) or  # Todo mayúsculas
            linea.startswith('#')  # Markdown header
        )):
            fin_idx = idx
            break

    # Extraer contenido
    contenido_seccion = '\n'.join(lineas[inicio_idx:fin_idx])

    # Formatear en markdown
    contenido_markdown = f"# {titulo_pagina}\n\n{contenido_seccion.strip()}"

    return contenido_markdown

def generar_resumen_documento(contenido_documento, filename):
    """
    Genera un resumen ejecutivo del documento usando Frida
    """
    prompt_resumen = f"""Genera un resumen ejecutivo del siguiente documento funcional.

**Documento:** {filename}

**Contenido:**
{contenido_documento[:10000]}

**Tu tarea:**
Crea un resumen ejecutivo de 300-500 palabras que incluya:
1. Objetivo principal del documento
2. Alcance del proyecto/sistema
3. Puntos clave y características principales
4. Enlaces a secciones detalladas

**Formato:** Markdown, profesional y claro.
**Importante:** Solo el contenido markdown, sin explicaciones adicionales."""

    payload = {
        "model": st.session_state.model,
        "messages": [
            {"role": "system", "content": "Eres un experto en resumir documentación técnica."},
            {"role": "user", "content": prompt_resumen}
        ],
        "temperature": 0.5
    }

    try:
        respuesta = call_ia(payload)
        return f"# Resumen General\n\n{respuesta.strip()}"
    except:
        return f"# Resumen General\n\n**Documento:** {filename}\n\n[Resumen pendiente de generar]"

def generar_glosario_documento(contenido_documento):
    """
    Genera un glosario de términos técnicos del documento usando Frida
    """
    prompt_glosario = f"""Identifica los términos técnicos, acrónimos y conceptos clave del siguiente documento y crea un glosario.

**Contenido:**
{contenido_documento[:8000]}

**Tu tarea:**
Crea un glosario en formato markdown con:
- Términos técnicos y acrónimos encontrados
- Definición clara de cada uno
- Orden alfabético

**Formato:** Markdown con lista de términos.
**Importante:** Solo el contenido markdown, sin explicaciones adicionales."""

    payload = {
        "model": st.session_state.model,
        "messages": [
            {"role": "system", "content": "Eres un experto en documentación técnica."},
            {"role": "user", "content": prompt_glosario}
        ],
        "temperature": 0.3
    }

    try:
        respuesta = call_ia(payload)
        return f"# Glosario\n\n{respuesta.strip()}"
    except:
        return "# Glosario\n\n[Glosario pendiente de generar]"

def analizar_documento_con_frida(contenido_documento, filename):
    """
    Usa Frida para analizar el documento y proponer una estructura de wiki
    Retorna una estructura jerárquica de páginas sugeridas
    """
    prompt_analisis = f"""Analiza el siguiente documento funcional y propón una estructura jerárquica de páginas Wiki.

**Documento:** {filename}

**Contenido del documento:**
{contenido_documento[:20000]}
{"[Documento truncado para análisis, pero el contenido COMPLETO se usará en la wiki...]" if len(contenido_documento) > 20000 else ""}

**IMPORTANTE:**
- El contenido COMPLETO del documento se incluirá en las páginas
- NO resumas ni omitas nada
- Solo propón la ESTRUCTURA (títulos y organización)

**Tu tarea:**
Propón una estructura jerárquica indicando para cada página:
1. Título de la página
2. Encabezados del documento original que corresponden a esta sección
3. Tipo de página: "resumen", "contenido_completo", o "glosario"

**Páginas especiales:**
- Página "Resumen General" (resumen ejecutivo corto)
- Página "Glosario" (solo si hay términos técnicos)

**Resto de páginas:**
- Deben contener el contenido COMPLETO de cada sección del documento
- Divide en secciones lógicas según los encabezados del documento

**Formato de respuesta (JSON):**

```json
{{
  "paginas": [
    {{
      "titulo": "Resumen General",
      "es_raiz": true,
      "padre": null,
      "tipo": "resumen",
      "seccion_origen": "",
      "orden": 0
    }},
    {{
      "titulo": "Introducción",
      "es_raiz": false,
      "padre": "Resumen General",
      "tipo": "contenido_completo",
      "seccion_origen": "1. Introducción|INTRODUCCIÓN|Introducción",
      "orden": 1
    }},
    {{
      "titulo": "Glosario",
      "es_raiz": false,
      "padre": "Resumen General",
      "tipo": "glosario",
      "seccion_origen": "",
      "orden": 99
    }}
  ]
}}
```

**Campos:**
- `tipo`: "resumen", "contenido_completo", o "glosario"
- `seccion_origen`: Encabezados del documento que corresponden (separados por |)

**Reglas:**
- Máximo 12 páginas de contenido + resumen + glosario
- Jerarquía de máximo 2 niveles
- Sigue el orden lógico del documento"""

    payload = {
        "model": st.session_state.model,
        "messages": [
            {
                "role": "system",
                "content": "Eres un experto en estructuración de documentación. Tu trabajo es proponer la ESTRUCTURA, no generar contenido."
            },
            {
                "role": "user",
                "content": prompt_analisis
            }
        ],
        "temperature": 0.3
    }

    try:
        with st.spinner("🧠 Paso 1/2: Frida está analizando la estructura del documento..."):
            respuesta = call_ia(payload)

        # Extraer JSON de la respuesta
        json_match = re.search(r'```json\s*(.*?)\s*```', respuesta, re.DOTALL)
        if json_match:
            json_str = json_match.group(1)
            estructura = json.loads(json_str)
        else:
            estructura = json.loads(respuesta)

        # Paso 2: Generar contenido completo para cada página
        total_paginas = len(estructura['paginas'])
        with st.spinner(f"📝 Paso 2/2: Extrayendo contenido completo ({total_paginas} páginas)..."):
            for idx, pagina in enumerate(estructura['paginas']):
                progress_msg = f"  [{idx+1}/{total_paginas}] {pagina['titulo']}"
                st.caption(progress_msg)

                if pagina['tipo'] == 'resumen':
                    pagina['contenido_markdown'] = generar_resumen_documento(contenido_documento, filename)
                elif pagina['tipo'] == 'glosario':
                    pagina['contenido_markdown'] = generar_glosario_documento(contenido_documento)
                else:  # contenido_completo
                    pagina['contenido_markdown'] = extraer_contenido_seccion(
                        contenido_documento,
                        pagina.get('seccion_origen', ''),
                        pagina['titulo']
                    )

        st.success(f"✅ Estructura generada con {total_paginas} páginas (contenido completo incluido)")
        return estructura

    except json.JSONDecodeError as e:
        st.error(f"Error al parsear respuesta de Frida: {str(e)}")
        st.code(respuesta[:1000])
        return None
    except Exception as e:
        st.error(f"Error al analizar documento: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return None

def mejorar_contenido_pagina_con_frida(titulo_pagina, contenido_original, contexto_documento=""):
    """
    Usa Frida para mejorar el contenido de una página específica
    """
    prompt_mejora = f"""Mejora el siguiente contenido para una página de Wiki en Azure DevOps.

**Título de la página:** {titulo_pagina}

**Contenido original:**
{contenido_original}

**Contexto del documento completo (si aplica):**
{contexto_documento[:3000] if contexto_documento else "No disponible"}

**Tu tarea:**
1. Reformula el contenido para mayor claridad
2. Estructura la información de forma lógica usando markdown
3. Añade ejemplos o aclaraciones donde sea útil
4. Mantén un tono profesional pero accesible
5. Usa listas, tablas, y formato markdown adecuadamente

**Devuelve solo el contenido markdown mejorado, sin explicaciones adicionales.**"""

    payload = {
        "model": st.session_state.model,
        "messages": [
            {
                "role": "system",
                "content": "Eres un experto en redacción técnica y documentación clara."
            },
            {
                "role": "user",
                "content": prompt_mejora
            }
        ],
        "temperature": 0.5
    }

    try:
        respuesta = call_ia(payload)
        return respuesta.strip()
    except Exception as e:
        st.error(f"Error al mejorar contenido: {str(e)}")
        return contenido_original

def crear_pagina_wiki_azure(organization, project, pat, wiki_id, path, contenido_markdown):
    """
    Crea o actualiza una página en Azure DevOps Wiki

    Parameters:
    - path: Ruta de la página (ej: "/Introduccion" o "/Introduccion/Objetivos")
    - contenido_markdown: Contenido en formato markdown

    Returns:
    - True si se creó exitosamente, False si hubo error
    """
    # Limpiar y formatear el path
    path = path.strip()
    if not path.startswith('/'):
        path = '/' + path

    # URL de la API
    url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki_id}/pages?path={path}&api-version=7.1"

    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Basic {encoded_credentials}"
    }

    payload = {
        "content": contenido_markdown
    }

    try:
        response = requests.put(url, json=payload, headers=headers, timeout=30)

        if response.status_code in [200, 201]:
            return True, response.json()
        elif response.status_code == 409:
            # La página ya existe, intentar actualizar
            return actualizar_pagina_wiki_azure(organization, project, pat, wiki_id, path, contenido_markdown)
        else:
            st.error(f"❌ Error {response.status_code} al crear página: {path}")
            st.code(response.text[:300])
            return False, None

    except Exception as e:
        st.error(f"❌ Error al crear página {path}: {str(e)}")
        return False, None

def actualizar_pagina_wiki_azure(organization, project, pat, wiki_id, path, contenido_markdown):
    """
    Actualiza una página existente en Azure DevOps Wiki
    """
    # Primero obtener la versión actual de la página
    url_get = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki_id}/pages?path={path}&api-version=7.1"

    credentials = f":{pat}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Basic {encoded_credentials}"
    }

    try:
        # Obtener ETag para actualización
        response_get = requests.get(url_get, headers=headers, timeout=30)

        if response_get.status_code != 200:
            st.error(f"No se pudo obtener información de la página: {path}")
            return False, None

        page_data = response_get.json()
        etag = response_get.headers.get('ETag')

        # Actualizar con el ETag
        headers['If-Match'] = etag

        payload = {
            "content": contenido_markdown
        }

        response_put = requests.put(url_get, json=payload, headers=headers, timeout=30)

        if response_put.status_code in [200, 201]:
            return True, response_put.json()
        else:
            st.error(f"❌ Error {response_put.status_code} al actualizar página: {path}")
            st.code(response_put.text[:300])
            return False, None

    except Exception as e:
        st.error(f"❌ Error al actualizar página {path}: {str(e)}")
        return False, None

def obtener_estructura_paginas_wiki_existente(organization, project, pat, wiki_id):
    """
    Obtiene la estructura de páginas existentes en la wiki para mostrar al usuario
    """
    paginas = obtener_paginas_wiki(organization, project, pat, wiki_id, recursion_level=5)

    if not paginas:
        return []

    # Crear estructura jerárquica para visualización
    estructura = []
    for pagina in paginas:
        path_parts = pagina['path'].strip('/').split('/')
        nivel = len(path_parts) - 1
        estructura.append({
            'path': pagina['path'],
            'id': pagina['id'],
            'nivel': nivel,
            'nombre': path_parts[-1] if path_parts else pagina['path']
        })

    return estructura

# ==================================================
# SIDEBAR
# ==================================================
st.sidebar.title("⚙️ Configuración")

if st.sidebar.button("🧹 Nuevo Chat"):
    for k in defaults:
        st.session_state[k] = defaults[k]

with st.sidebar.expander("⚙️ Configuración IA", expanded=False):
    st.session_state.model = st.selectbox(
        "Modelo IA",
        options=[
            "Innovation-gpt4o-mini", "Innovation-gpt4o",
            "o4-mini", "o1", "o1-mini", "o3-mini",
            "o1-preview", "gpt-5-chat", "gpt-4.1",
            "gpt-4.1-mini", "gpt-5", "gpt-5-codex",
            "gpt-5-mini", "gpt-5-nano",
            "gpt-4.1-nano", "claude-3-5-sonnet",
            "claude-4-sonnet", "claude-3-7-sonnet",
            "claude-3-5-haiku", "claude-4-5-sonnet"
        ],
        index=0
    )
    st.session_state.include_temp = st.checkbox("Incluir temperatura", value=True)
    st.session_state.temperature = st.slider("Temperatura", 0.0, 1.0, 0.7, 0.1)
    st.session_state.include_tokens = st.checkbox("Incluir max_tokens", value=True)
    st.session_state.max_tokens = st.slider("Max tokens", 100, 4096, 3000, 100)

with st.sidebar.expander("📝 Config Prompt Chat", expanded=False):
    template_type = st.selectbox(
        "Tipo de prompt inicial",
        [
            "Libre", "PO Casos exito", "PO Definicion epica",
            "PO Definicion epica una historia", "PO Definicion historia",
            "PO Definicion mejora tecnica", "PO Definicion spike",
            "PO resumen reunion", "Programador Python"
        ]
    )
    prompt_template = st.text_area("Contenido del template", get_template(template_type), height=220)

# Configuración de Azure DevOps (global)
st.sidebar.markdown("---")
st.sidebar.markdown("### 🔗 Azure DevOps")

with st.sidebar.expander("Configurar conexión", expanded=False):
    org_input = st.text_input(
        "Organización",
        value=st.session_state.devops_org if st.session_state.devops_org else "TelepizzaIT",
        placeholder="ej: TelepizzaIT",
        key="sidebar_org_input"
    )
    project_input = st.text_input(
        "Proyecto",
        value=st.session_state.devops_project if st.session_state.devops_project else "Sales",
        placeholder="ej: Sales",
        key="sidebar_project_input"
    )
    pat_input = st.text_input(
        "PAT",
        value=st.session_state.devops_pat,
        type="password",
        help="Personal Access Token",
        key="sidebar_pat_input"
    )

    if st.button("💾 Guardar", use_container_width=True, key="sidebar_save_devops"):
        if org_input and project_input and pat_input:
            st.session_state.devops_org = org_input
            st.session_state.devops_project = project_input
            st.session_state.devops_pat = pat_input
            st.success("✅ Guardado")
            st.rerun()
        else:
            st.error("❌ Completa todos los campos")

if st.session_state.devops_org and st.session_state.devops_project and st.session_state.devops_pat:
    st.sidebar.success(f"✅ Conectado: {st.session_state.devops_org}/{st.session_state.devops_project}")
else:
    st.sidebar.info("ℹ️ Configura Azure DevOps para usar todas las funcionalidades")

# ==================================================
# TABS
# ==================================================
tab_chat, tab_devops, tab_doc = st.tabs([
    "💬 Chat clásico",
    "🎯 Tareas Azure DevOps",
    "📄 Análisis Documentos"
])

# ================= TAB 1: CHAT CLÁSICO =================
with tab_chat:
    st.title("💬 Chat Softtek Prompts IA")
    for m in st.session_state.messages:
        with st.chat_message(m["role"]):
            st.markdown(m["content"])
    if prompt := st.chat_input("Escribe tu mensaje..."):
        prompt_final = prompt_template.format(input=prompt) if not st.session_state.messages else prompt
        st.session_state.messages.append({"role": "user", "content": prompt_final})
        payload = {"model": st.session_state.model, "messages": st.session_state.messages}
        if st.session_state.include_temp:
            payload["temperature"] = st.session_state.temperature
        if st.session_state.include_tokens:
            payload["max_tokens"] = st.session_state.max_tokens
        with st.spinner("🤖 La IA está pensando..."):
            answer = call_ia(payload)
        st.session_state.messages.append({"role": "assistant", "content": answer})
        st.rerun()

# ================= TAB 2: CONSULTA TAREAS DEVOPS =================
with tab_devops:
    st.title("🎯 Azure DevOps")
    st.markdown("Consulta work items y documentación Wiki de Azure DevOps usando IA")

    # Verificar conexión
    if not st.session_state.devops_org or not st.session_state.devops_project or not st.session_state.devops_pat:
        st.info("ℹ️ **Configura Azure DevOps en el sidebar** (⚙️ Configuración → 🔗 Azure DevOps)")
        st.markdown("""
        **Permisos necesarios del PAT:**
        - Work Items (Read)
        - Wiki (Read) o Code (Read)

        [¿Cómo crear un PAT?](https://docs.microsoft.com/en-us/azure/devops/organizations/accounts/use-personal-access-tokens-to-authenticate)
        """)

    st.markdown("---")

    # Crear subtabs
    subtab_workitems, subtab_wiki, subtab_crear_wiki, subtab_crear_tarea = st.tabs([
        "📋 Consulta Work Items",
        "📚 Consulta Wiki",
        "🔨 Crear Wiki desde Documento",
        "➕ Crear Tarea"
    ])

    # ================= SUBTAB 1: CONSULTA WORK ITEMS =================
    with subtab_workitems:
        # Verificar conexión
        if not st.session_state.devops_pat or not st.session_state.devops_org or not st.session_state.devops_project:
            st.warning("⚠️ Primero configura la conexión a Azure DevOps en la sección de Configuración arriba")
        else:
            # Configuración de sincronización
            with st.expander("🎛️ Filtros y Sincronización de Work Items", expanded=not st.session_state.devops_indexed):
                col_filtros1, col_filtros2 = st.columns(2)
    
                with col_filtros1:
                    work_item_types = st.multiselect(
                        "Tipos de Work Items",
                        options=["Bug", "User Story", "Task", "Feature", "Epic", "Issue", "Test Case"],
                        default=["Bug"],
                        help="Selecciona uno o más tipos"
                    )
    
                    area_path_input = st.text_input(
                        "Área (opcional)",
                        value="",
                        placeholder="ej: Sales\\MySaga POC",
                        help="Deja vacío para todas las áreas"
                    )
    
                with col_filtros2:
                    max_items = st.slider(
                        "Límite de items a traer",
                        min_value=50,
                        max_value=1000,
                        value=200,
                        step=50,
                        help="Máximo de work items a sincronizar"
                    )
    
                    top_k_similar = st.slider(
                        "Items similares a mostrar",
                        min_value=3,
                        max_value=10,
                        value=5,
                        step=1,
                        help="Número de items similares para enviar a Frida"
                    )
    
                st.markdown("---")
    
                col_btn1, col_btn2 = st.columns([3, 1])
                with col_btn1:
                    if st.button("🔄 Sincronizar e Indexar Work Items", use_container_width=True):
                        if not work_item_types or len(work_item_types) == 0:
                            st.error("❌ Selecciona al menos un tipo de work item")
                        else:
                            with st.spinner("📥 Obteniendo work items de Azure DevOps..."):
                                incidencias = obtener_incidencias_devops(
                                    st.session_state.devops_org,
                                    st.session_state.devops_project,
                                    st.session_state.devops_pat,
                                    area_path=area_path_input if area_path_input else None,
                                    work_item_types=work_item_types,
                                    max_items=max_items
                                )
    
                            if incidencias:
                                st.success(f"✅ Se encontraron {len(incidencias)} work items")
    
                                tipos_count = {}
                                for inc in incidencias:
                                    tipo = inc['tipo']
                                    tipos_count[tipo] = tipos_count.get(tipo, 0) + 1
    
                                st.info(f"📊 Distribución: " + ", ".join([f"{t}: {c}" for t, c in tipos_count.items()]))
    
                                st.session_state.devops_incidencias = incidencias
    
                                if st.session_state.embedding_model is None:
                                    st.session_state.embedding_model = cargar_modelo_embeddings()
    
                                embeddings = generar_embeddings_incidencias(
                                    incidencias,
                                    st.session_state.embedding_model
                                )
                                st.session_state.devops_embeddings = embeddings
                                st.session_state.devops_indexed = True
                                st.session_state.devops_top_k = top_k_similar
    
                                st.success("✅ Indexación completada. Ahora puedes hacer consultas.")
                                st.rerun()
                            else:
                                st.warning("⚠️ No se encontraron work items o hubo un error")
    
                with col_btn2:
                    if st.button("🗑️ Limpiar", use_container_width=True, key="limpiar_devops"):
                        st.session_state.devops_incidencias = []
                        st.session_state.devops_embeddings = None
                        st.session_state.devops_indexed = False
                        st.session_state.devops_messages = []
                        st.success("✅ Cache limpiado")
                        st.rerun()
    
            # Estado de indexación
            if st.session_state.devops_indexed:
                tipos_en_cache = {}
                for inc in st.session_state.devops_incidencias:
                    tipo = inc['tipo']
                    tipos_en_cache[tipo] = tipos_en_cache.get(tipo, 0) + 1
    
                tipos_str = ", ".join([f"{t} ({c})" for t, c in tipos_en_cache.items()])
                st.info(f"📊 **{len(st.session_state.devops_incidencias)} work items indexados**: {tipos_str}")
                st.info(f"🎯 **Top-K configurado**: {st.session_state.get('devops_top_k', 5)} items similares por consulta")
    
            # Chat de consultas
            st.markdown("---")
    
            col_chat, col_stats = st.columns([2, 1])
    
            with col_stats:
                st.subheader("📈 Estadísticas")
                if st.session_state.devops_incidencias:
                    st.markdown("**Por tipo:**")
                    tipos = {}
                    for inc in st.session_state.devops_incidencias:
                        tipo = inc['tipo']
                        tipos[tipo] = tipos.get(tipo, 0) + 1
                    for tipo, count in sorted(tipos.items()):
                        st.metric(tipo, count)
    
                    st.markdown("---")
    
                    st.markdown("**Por estado:**")
                    estados = {}
                    for inc in st.session_state.devops_incidencias:
                        estado = inc['estado']
                        estados[estado] = estados.get(estado, 0) + 1
                    for estado, count in sorted(estados.items(), key=lambda x: x[1], reverse=True)[:5]:
                        st.text(f"{estado}: {count}")
                else:
                    st.info("Sincroniza primero los work items")
    
            with col_chat:
                st.subheader("💬 Chat de consultas")
    
                for m in st.session_state.devops_messages:
                    with st.chat_message(m["role"]):
                        st.markdown(m["content"])
    
                if devops_query := st.chat_input(
                    "Pregunta sobre work items... ej: '¿Cómo se implementó X?'",
                    key="devops_chat",
                    disabled=not st.session_state.devops_indexed
                ):
                    if not st.session_state.devops_indexed:
                        st.warning("⚠️ Primero debes sincronizar e indexar los work items")
                    else:
                        st.session_state.devops_messages.append({"role": "user", "content": devops_query})
    
                        top_k = st.session_state.get('devops_top_k', 5)
    
                        with st.spinner(f"🔍 Buscando los {top_k} work items más similares..."):
                            resultados = buscar_incidencias_similares(
                                devops_query,
                                st.session_state.devops_incidencias,
                                st.session_state.devops_embeddings,
                                st.session_state.embedding_model,
                                top_k=top_k
                            )
    
                        contexto = construir_contexto_devops(resultados)
    
                        system_prompt = """Eres un asistente técnico experto en analizar work items de software.
    
    Cuando respondas:
    1. Analiza los work items similares (pueden ser Bugs, User Stories, Tasks, etc.)
    2. Si hay coincidencia exacta o similar, explica cómo se resolvió o implementó
    3. Si no hay coincidencia exacta, propón soluciones basadas en casos similares
    4. Sé específico y técnico
    5. Menciona el ID y tipo de los work items relevantes
    6. Si encuentras patrones comunes, menciónalo"""
    
                        payload = {
                            "model": st.session_state.model,
                            "messages": [
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": f"{contexto}\n\n**Consulta:** {devops_query}"}
                            ]
                        }
    
                        if st.session_state.include_temp:
                            payload["temperature"] = st.session_state.temperature
                        if st.session_state.include_tokens:
                            payload["max_tokens"] = st.session_state.max_tokens
    
                        with st.spinner("🤖 Frida está analizando los work items..."):
                            respuesta = call_ia(payload)
    
                        st.session_state.devops_messages.append({"role": "assistant", "content": respuesta})
    
                        with st.expander(f"📋 Ver los {top_k} work items más similares"):
                            for i, resultado in enumerate(resultados, 1):
                                inc = resultado["incidencia"]
                                sim = resultado["similitud"]
    
                                st.markdown(f"### Work Item {i} - [{inc['tipo']}] ID: {inc['id']} (Similitud: {sim:.1%})")
                                st.markdown(f"**Título:** {inc['titulo']}")
                                st.markdown(f"**Estado:** {inc['estado']}")
                                st.markdown(f"**Área:** {inc['area']}")
                                if inc['resolucion']:
                                    st.markdown(f"**Resolución:** {inc['resolucion']}")
                                st.markdown(f"**Descripción:** {limpiar_html(inc['descripcion'])[:300]}...")
                                if inc['tags']:
                                    st.markdown(f"**Tags:** {inc['tags']}")
                                st.markdown("---")
    
                        st.rerun()

    # ================= SUBTAB 2: CONSULTA WIKI =================
    with subtab_wiki:
        st.subheader("📚 Consulta Wiki de Azure DevOps")
        st.markdown("Indexa páginas de la Wiki y haz consultas sobre su contenido")

        # Verificar configuración de Azure DevOps
        if not st.session_state.devops_pat or not st.session_state.devops_org or not st.session_state.devops_project:
            st.warning("⚠️ Primero configura la conexión a Azure DevOps en la sección de Configuración arriba")
        else:
            # Sección de carga de Wiki
            with st.expander("📥 Seleccionar Páginas de Wiki", expanded=not st.session_state.wiki_indexed):
                col1, col2 = st.columns([2, 1])
    
                with col1:
                    # Listar wikis disponibles
                    if st.button("🔍 Listar Wikis del Proyecto"):
                        with st.spinner("Obteniendo wikis..."):
                            wikis = obtener_wikis_proyecto(
                                st.session_state.devops_org,
                                st.session_state.devops_project,
                                st.session_state.devops_pat
                            )
    
                        if wikis:
                            st.session_state.available_wikis = wikis
                            st.success(f"✅ {len(wikis)} wiki(s) encontrada(s)")
                        else:
                            st.error("❌ No se encontraron wikis o hubo un error")
    
                    # Selector de wiki
                    if 'available_wikis' in st.session_state and st.session_state.available_wikis:
                        selected_wiki_idx = st.selectbox(
                            "Selecciona una Wiki",
                            options=range(len(st.session_state.available_wikis)),
                            format_func=lambda i: f"{st.session_state.available_wikis[i]['name']} ({st.session_state.available_wikis[i]['type']})",
                            key="wiki_selector"
                        )
    
                        selected_wiki = st.session_state.available_wikis[selected_wiki_idx]
                        st.session_state.selected_wiki_id = selected_wiki['id']
                        st.session_state.selected_wiki_name = selected_wiki['name']
    
                        st.info(f"📖 Wiki seleccionada: **{selected_wiki['name']}**")
    
                        # Botón para listar páginas
                        if st.button("📄 Listar Páginas de esta Wiki"):
                            # Guardar logs en session_state para mostrar en col2
                            st.session_state.wiki_logs = []
    
                            with st.spinner("Obteniendo páginas principales..."):
                                paginas = obtener_paginas_wiki(
                                    st.session_state.devops_org,
                                    st.session_state.devops_project,
                                    st.session_state.devops_pat,
                                    st.session_state.selected_wiki_id
                                )
    
                            if paginas:
                                # Obtener subpáginas para cada página padre
                                todas_paginas = []
                                total_principales = len(paginas)
                                st.session_state.wiki_logs.append(("info", f"📊 Páginas principales: {total_principales}"))
    
                                for idx, pagina in enumerate(paginas):
                                    todas_paginas.append(pagina)
    
                                    # Si es página padre, obtener sus subpáginas
                                    if pagina.get('isParentPage', False):
                                        with st.spinner(f"📁 Obteniendo subpáginas de {pagina['path']} ({idx+1}/{total_principales})..."):
                                            subpaginas = obtener_subpaginas_especificas(
                                                st.session_state.devops_org,
                                                st.session_state.devops_project,
                                                st.session_state.devops_pat,
                                                st.session_state.selected_wiki_id,
                                                pagina['path']
                                            )
    
                                            if subpaginas:
                                                st.session_state.wiki_logs.append(("info", f"  └─ {pagina['path']}: +{len(subpaginas)} subpágina(s)"))
                                                todas_paginas.extend(subpaginas)
    
                                st.session_state.available_wiki_pages = todas_paginas
                                st.session_state.wiki_logs.append(("success", f"✅ Total: {len(todas_paginas)} página(s) ({total_principales} principales + {len(todas_paginas) - total_principales} subpáginas)"))
                                st.rerun()
                            else:
                                st.session_state.wiki_logs.append(("warning", "⚠️ No se encontraron páginas en esta wiki"))
                                st.rerun()
    
                        # Selector de páginas (individual + batch)
                        if 'available_wiki_pages' in st.session_state and st.session_state.available_wiki_pages:
                            st.markdown("#### Seleccionar páginas para indexar:")
    
                            # Opción: Seleccionar todas
                            select_all = st.checkbox("Seleccionar todas las páginas", value=False)
    
                            # Lista de checkboxes para páginas
                            selected_pages = []
    
                            if select_all:
                                selected_pages = st.session_state.available_wiki_pages.copy()
                                st.info(f"📑 Todas las páginas seleccionadas ({len(selected_pages)})")
                            else:
                                st.markdown("**Selecciona páginas individualmente:**")
                                for idx, page in enumerate(st.session_state.available_wiki_pages):
                                    if st.checkbox(
                                        f"{page['path']}",
                                        value=False,
                                        key=f"wiki_page_{idx}"
                                    ):
                                        selected_pages.append(page)
    
                            st.session_state.selected_wiki_pages = selected_pages
    
                            if selected_pages:
                                st.success(f"✅ {len(selected_pages)} página(s) seleccionada(s)")
    
                with col2:
                    st.markdown("#### ⚙️ Configuración")
                    wiki_chunk_size = st.slider(
                        "Tamaño de fragmentos",
                        min_value=500,
                        max_value=4000,
                        value=1000,
                        step=100,
                        help="Tamaño de cada fragmento de las páginas Wiki"
                    )
    
                    wiki_top_k = st.slider(
                        "Fragmentos relevantes",
                        min_value=3,
                        max_value=10,
                        value=5,
                        step=1,
                        help="Número de fragmentos a usar como contexto"
                    )
    
                    st.markdown("---")
                    st.markdown("#### 📊 Logs y Debug")
    
                    # Mostrar logs guardados
                    if hasattr(st.session_state, 'wiki_logs') and st.session_state.wiki_logs:
                        for log_type, log_message in st.session_state.wiki_logs:
                            if log_type == "info":
                                st.info(log_message)
                            elif log_type == "success":
                                st.success(log_message)
                            elif log_type == "warning":
                                st.warning(log_message)
                            elif log_type == "error":
                                st.error(log_message)
                    else:
                        st.text("No hay logs disponibles")
    
                st.markdown("---")
    
                # Botones de acción
                col_btn1, col_btn2 = st.columns([3, 1])
    
                with col_btn1:
                    if st.button("🔄 Procesar e Indexar Páginas", use_container_width=True, key="procesar_wiki_btn"):
                        if not hasattr(st.session_state, 'selected_wiki_pages') or len(st.session_state.selected_wiki_pages) == 0:
                            st.error("❌ Debes seleccionar al menos una página de la wiki")
                        else:
                            # Procesar cada página seleccionada
                            paginas_contenido = []
                            progress_bar = st.progress(0)
                            total_pages = len(st.session_state.selected_wiki_pages)
    
                            for idx, page in enumerate(st.session_state.selected_wiki_pages):
                                progress_bar.progress((idx + 1) / total_pages)
    
                                with st.spinner(f"Descargando {page['path']}..."):
                                    contenido_page = obtener_contenido_pagina_wiki(
                                        st.session_state.devops_org,
                                        st.session_state.devops_project,
                                        st.session_state.devops_pat,
                                        st.session_state.selected_wiki_id,
                                        page['id']
                                    )
    
                                if contenido_page and contenido_page['content']:
                                    # Limpiar markdown
                                    texto_limpio = limpiar_markdown(contenido_page['content'])
    
                                    # Dividir en chunks
                                    chunks = dividir_en_chunks(texto_limpio, chunk_size=wiki_chunk_size)
    
                                    paginas_contenido.append({
                                        'id': page['id'],
                                        'path': page['path'],
                                        'chunks': chunks
                                    })
    
                            if paginas_contenido:
                                st.success(f"✅ {len(paginas_contenido)} página(s) procesada(s)")
    
                                # Cargar modelo si no está cargado
                                if st.session_state.embedding_model is None:
                                    st.session_state.embedding_model = cargar_modelo_embeddings()
    
                                # Generar embeddings
                                embeddings, referencias = generar_embeddings_wiki(
                                    paginas_contenido,
                                    st.session_state.embedding_model
                                )
    
                                # Extraer lista plana de chunks para búsqueda
                                todos_chunks = []
                                for pagina in paginas_contenido:
                                    todos_chunks.extend(pagina['chunks'])
    
                                # Guardar en session_state
                                st.session_state.wiki_paginas_contenido = paginas_contenido
                                st.session_state.wiki_embeddings = embeddings
                                st.session_state.wiki_referencias = referencias
                                st.session_state.wiki_chunks = todos_chunks
                                st.session_state.wiki_indexed = True
                                st.session_state.wiki_top_k = wiki_top_k
    
                                total_chunks = sum(len(p['chunks']) for p in paginas_contenido)
                                st.success(f"✅ Wiki indexada: {len(paginas_contenido)} páginas, {total_chunks} fragmentos")
                                st.rerun()
                            else:
                                st.error("❌ No se pudo procesar ninguna página")
    
                with col_btn2:
                    if st.button("🗑️ Limpiar", use_container_width=True, key="limpiar_wiki"):
                        st.session_state.wiki_paginas_contenido = []
                        st.session_state.wiki_embeddings = None
                        st.session_state.wiki_referencias = []
                        st.session_state.wiki_chunks = []
                        st.session_state.wiki_indexed = False
                        st.session_state.wiki_messages = []
                        if 'available_wikis' in st.session_state:
                            del st.session_state.available_wikis
                        if 'available_wiki_pages' in st.session_state:
                            del st.session_state.available_wiki_pages
                        if 'selected_wiki_pages' in st.session_state:
                            del st.session_state.selected_wiki_pages
                        st.success("✅ Wiki limpiada")
                        st.rerun()
    
            # Estado de indexación
            if st.session_state.wiki_indexed:
                total_chunks = len(st.session_state.wiki_chunks)
                total_pages = len(st.session_state.wiki_paginas_contenido)
                st.info(f"📚 **Wiki indexada**: {st.session_state.selected_wiki_name} - {total_pages} páginas, {total_chunks} fragmentos")
                st.info(f"🎯 **Top-K configurado**: {st.session_state.get('wiki_top_k', 5)} fragmentos por consulta")
    
            st.markdown("---")
    
            # Chat de consultas Wiki
            st.subheader("💬 Consultas sobre la Wiki")
    
            for m in st.session_state.wiki_messages:
                with st.chat_message(m["role"]):
                    st.markdown(m["content"])
    
            if wiki_query := st.chat_input(
                "Pregunta sobre la Wiki... ej: '¿Cómo configurar X?'",
                key="wiki_chat",
                disabled=not st.session_state.wiki_indexed
            ):
                if not st.session_state.wiki_indexed:
                    st.warning("⚠️ Primero debes seleccionar e indexar páginas de la Wiki")
                else:
                    st.session_state.wiki_messages.append({"role": "user", "content": wiki_query})
    
                    top_k = st.session_state.get('wiki_top_k', 5)
    
                    # Buscar chunks relevantes
                    with st.spinner(f"🔍 Buscando fragmentos relevantes en la Wiki..."):
                        resultados = buscar_chunks_wiki_similares(
                            wiki_query,
                            st.session_state.wiki_chunks,
                            st.session_state.wiki_embeddings,
                            st.session_state.wiki_referencias,
                            st.session_state.embedding_model,
                            top_k=top_k
                        )
    
                    # Construir contexto
                    contexto = construir_contexto_wiki(resultados)
    
                    # Llamar a Frida
                    system_prompt = """Eres un asistente experto en analizar documentación técnica de wikis.
    
    Cuando respondas:
    1. Basa tu respuesta en la información de los fragmentos de la Wiki proporcionados
    2. Si la información no está en los fragmentos, indícalo claramente
    3. Menciona las páginas específicas de la Wiki cuando sea relevante
    4. Proporciona respuestas claras y estructuradas
    5. Si encuentras procedimientos o pasos, enuméralos claramente"""
    
                    payload = {
                        "model": st.session_state.model,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"{contexto}\n\n**Consulta:** {wiki_query}"}
                        ]
                    }
    
                    if st.session_state.include_temp:
                        payload["temperature"] = st.session_state.temperature
                    if st.session_state.include_tokens:
                        payload["max_tokens"] = st.session_state.max_tokens
    
                    with st.spinner("🤖 Frida está analizando la Wiki..."):
                        respuesta = call_ia(payload)
    
                    st.session_state.wiki_messages.append({"role": "assistant", "content": respuesta})
    
                    # Mostrar fragmentos usados
                    with st.expander(f"📋 Ver fragmentos de Wiki utilizados"):
                        for i, resultado in enumerate(resultados, 1):
                            sim = resultado["similitud"]
                            chunk = resultado["chunk"]
                            path = resultado["path"]
    
                            st.markdown(f"### Fragmento {i} - Página: {path}")
                            st.markdown(f"**Relevancia:** {sim:.1%}")
                            st.text(chunk[:500] + ("..." if len(chunk) > 500 else ""))
                            st.markdown("---")
    
                    st.rerun()
    
    # ================= SUBTAB 3: CREAR WIKI DESDE DOCUMENTO =================
    with subtab_crear_wiki:
        st.subheader("🔨 Crear Estructura Wiki desde Documento Funcional")
        st.markdown("Carga un documento funcional (.docx o .pdf), analízalo con Frida, y crea páginas Wiki en Azure DevOps")

        # Verificar configuración de Azure DevOps
        if not st.session_state.devops_pat or not st.session_state.devops_org or not st.session_state.devops_project:
            st.warning("⚠️ Primero configura la conexión a Azure DevOps en la sección de Configuración arriba")
        else:
            # === PASO 1: CARGA DEL DOCUMENTO ===
            with st.expander("📥 Paso 1: Cargar Documento Funcional", expanded=(not st.session_state.wiki_create_doc_content)):
                st.markdown("**Sube un documento funcional (.docx o .pdf)**")
    
                col_upload, col_info = st.columns([2, 1])
    
                with col_upload:
                    uploaded_file = st.file_uploader(
                        "Selecciona documento",
                        type=["docx", "pdf"],
                        key="wiki_create_upload",
                        help="Documento Word o PDF con especificaciones funcionales"
                    )
    
                    if uploaded_file:
                        file_ext = uploaded_file.name.split('.')[-1].lower()
    
                        if st.button("📖 Procesar Documento", key="procesar_doc_wiki"):
                            file_bytes = uploaded_file.read()
    
                            with st.spinner(f"📖 Leyendo {file_ext.upper()}..."):
                                if file_ext == "docx":
                                    contenido = leer_docx_desde_bytes(file_bytes)
                                elif file_ext == "pdf":
                                    contenido = leer_pdf_desde_bytes(file_bytes)
                                else:
                                    st.error("Formato no soportado")
                                    contenido = ""
    
                            if contenido:
                                st.session_state.wiki_create_doc_content = contenido
                                st.session_state.wiki_create_doc_filename = uploaded_file.name
                                st.success(f"✅ Documento procesado: {len(contenido)} caracteres")
                                st.rerun()
                            else:
                                st.error("❌ No se pudo leer el documento")
    
                with col_info:
                    if st.session_state.wiki_create_doc_content:
                        st.metric("📄 Documento cargado", st.session_state.wiki_create_doc_filename)
                        st.metric("📝 Tamaño", f"{len(st.session_state.wiki_create_doc_content)} caracteres")
    
                        if st.button("🗑️ Limpiar documento", key="limpiar_doc_wiki_create"):
                            st.session_state.wiki_create_doc_content = ""
                            st.session_state.wiki_create_doc_filename = ""
                            st.session_state.wiki_create_estructura_propuesta = None
                            st.session_state.wiki_create_estructura_editada = None
                            st.session_state.wiki_create_ready_to_create = False
                            st.rerun()
    
            # === PASO 2: ANÁLISIS CON FRIDA O MODO SIMPLE ===
            if st.session_state.wiki_create_doc_content:
                with st.expander("🧠 Paso 2: Elegir Modo de Creación", expanded=(not st.session_state.wiki_create_estructura_propuesta)):
                    st.markdown("**Elige cómo quieres organizar el contenido en la wiki:**")
    
                    modo_creacion = st.radio(
                        "Modo",
                        options=["analisis", "simple_una_pagina", "simple_dos_paginas", "dividir_por_encabezados"],
                        format_func=lambda x: {
                            "analisis": "✨ Analizar con Frida (estructura inteligente)",
                            "simple_una_pagina": "📄 Documento completo en 1 página (literal, sin análisis)",
                            "simple_dos_paginas": "📑 2 páginas: Documento completo + Resumen",
                            "dividir_por_encabezados": "📑 Dividir por encabezados (1 subpágina por punto principal)"
                        }[x],
                        help="Selecciona cómo organizar el contenido"
                    )
    
                    st.markdown("---")
    
                    # MODO ANÁLISIS CON FRIDA
                    if modo_creacion == "analisis":
                        st.markdown("""
                        **Frida analizará el documento y propondrá:**
                        - Estructura jerárquica de páginas
                        - Página de resumen ejecutivo
                        - Contenido completo extraído por secciones
                        - Página de glosario (si aplica)
                        """)
    
                        col_analizar, col_reset = st.columns([3, 1])
    
                        with col_analizar:
                            if st.button("✨ Analizar con Frida", use_container_width=True, key="analizar_frida"):
                                estructura = analizar_documento_con_frida(
                                    st.session_state.wiki_create_doc_content,
                                    st.session_state.wiki_create_doc_filename
                                )
    
                                if estructura and "paginas" in estructura:
                                    st.session_state.wiki_create_estructura_propuesta = estructura
                                    st.session_state.wiki_create_estructura_editada = json.loads(json.dumps(estructura))
                                    st.success(f"✅ Estructura propuesta: {len(estructura['paginas'])} páginas")
                                    st.rerun()
                                else:
                                    st.error("❌ No se pudo generar estructura. Intenta de nuevo.")
    
                        with col_reset:
                            if st.session_state.wiki_create_estructura_propuesta:
                                if st.button("🔄 Re-analizar", use_container_width=True, key="reanalizar_frida"):
                                    st.session_state.wiki_create_estructura_propuesta = None
                                    st.session_state.wiki_create_estructura_editada = None
                                    st.rerun()
    
                    # MODO SIMPLE: UNA PÁGINA
                    elif modo_creacion == "simple_una_pagina":
                        st.markdown("""
                        **📄 Modo simple:**
                        - Todo el contenido del documento en una sola página
                        - Sin análisis, sin modificaciones
                        - Contenido literal del documento
                        """)
    
                        if st.button("📄 Crear Estructura Simple (1 página)", use_container_width=True, key="crear_simple_1"):
                            estructura = {
                                "paginas": [
                                    {
                                        "titulo": st.session_state.wiki_create_doc_filename.replace('.docx', '').replace('.pdf', ''),
                                        "es_raiz": True,
                                        "padre": None,
                                        "contenido_markdown": f"# {st.session_state.wiki_create_doc_filename}\n\n{st.session_state.wiki_create_doc_content}",
                                        "orden": 0
                                    }
                                ]
                            }
                            st.session_state.wiki_create_estructura_propuesta = estructura
                            st.session_state.wiki_create_estructura_editada = json.loads(json.dumps(estructura))
                            st.session_state.wiki_create_ready_to_create = True
                            st.success("✅ Estructura creada: 1 página con todo el contenido")
                            st.rerun()
    
                    # MODO SIMPLE: DOS PÁGINAS
                    elif modo_creacion == "simple_dos_paginas":
                        st.markdown("""
                        **📑 Modo 2 páginas:**
                        - Página 1: Resumen ejecutivo (generado por Frida)
                        - Página 2: Documento completo (contenido literal)
                        """)
    
                        if st.button("📑 Crear Estructura (2 páginas)", use_container_width=True, key="crear_simple_2"):
                            with st.spinner("Generando resumen..."):
                                resumen = generar_resumen_documento(
                                    st.session_state.wiki_create_doc_content,
                                    st.session_state.wiki_create_doc_filename
                                )
    
                            doc_titulo = st.session_state.wiki_create_doc_filename.replace('.docx', '').replace('.pdf', '')
                            estructura = {
                                "paginas": [
                                    {
                                        "titulo": "Resumen",
                                        "es_raiz": True,
                                        "padre": None,
                                        "contenido_markdown": resumen,
                                        "orden": 0
                                    },
                                    {
                                        "titulo": doc_titulo,
                                        "es_raiz": False,
                                        "padre": "Resumen",
                                        "contenido_markdown": f"# {doc_titulo}\n\n{st.session_state.wiki_create_doc_content}",
                                        "orden": 1
                                    }
                                ]
                            }
                            st.session_state.wiki_create_estructura_propuesta = estructura
                            st.session_state.wiki_create_estructura_editada = json.loads(json.dumps(estructura))
                            st.session_state.wiki_create_ready_to_create = True
                            st.success("✅ Estructura creada: 2 páginas (Resumen + Documento completo)")
                            st.rerun()
    
                    # MODO DIVIDIR POR ENCABEZADOS
                    elif modo_creacion == "dividir_por_encabezados":
                        st.markdown("""
                        **📑 Modo dividir por encabezados:**
                        - Detecta automáticamente los puntos principales del documento
                        - Crea 1 subpágina por cada encabezado encontrado
                        - Contenido literal de cada sección (sin análisis)
                        - Página índice automática
                        """)
    
                        if st.button("📑 Dividir Automáticamente", use_container_width=True, key="dividir_encabezados"):
                            with st.spinner("🔍 Detectando encabezados principales..."):
                                estructura = dividir_documento_por_encabezados(
                                    st.session_state.wiki_create_doc_content,
                                    st.session_state.wiki_create_doc_filename
                                )
    
                            if estructura and "paginas" in estructura:
                                num_paginas = len(estructura['paginas'])
                                st.session_state.wiki_create_estructura_propuesta = estructura
                                st.session_state.wiki_create_estructura_editada = json.loads(json.dumps(estructura))
                                st.session_state.wiki_create_ready_to_create = True
    
                                if num_paginas > 1:
                                    st.success(f"✅ Estructura creada: {num_paginas} páginas (1 índice + {num_paginas-1} secciones)")
                                else:
                                    st.warning("⚠️ No se detectaron encabezados suficientes. Se creó 1 página con todo el contenido.")
                                st.rerun()
    
            # === PASO 3: REVISAR Y EDITAR ESTRUCTURA ===
            if st.session_state.wiki_create_estructura_propuesta:
                with st.expander("📝 Paso 3: Revisar y Editar Estructura Propuesta", expanded=True):
                    st.markdown("**Revisa la estructura propuesta por Frida. Puedes editar el contenido de cada página.**")
    
                    estructura = st.session_state.wiki_create_estructura_editada
    
                    # Mostrar árbol de páginas
                    st.markdown("#### 🌳 Estructura de Páginas:")
    
                    for idx, pagina in enumerate(estructura['paginas']):
                        nivel = 0 if pagina.get('es_raiz', False) else 1
                        if pagina.get('padre') and not pagina.get('es_raiz', False):
                            # Contar niveles basándose en la jerarquía
                            padre = pagina.get('padre', '')
                            for p in estructura['paginas']:
                                if p['titulo'] == padre and not p.get('es_raiz', False):
                                    nivel = 2
                                    break
    
                        indent = "　　" * nivel
                        icon = "📄" if pagina.get('es_raiz', False) else ("📁" if nivel == 1 else "📃")
    
                        with st.container():
                            col_titulo, col_acciones = st.columns([4, 1])
    
                            with col_titulo:
                                st.markdown(f"{indent}{icon} **{pagina['titulo']}**")
    
                            with col_acciones:
                                if st.button("✏️ Editar", key=f"edit_page_{idx}"):
                                    st.session_state[f"editing_page_{idx}"] = True
                                    st.rerun()
    
                            # Mostrar editor si está en modo edición
                            if st.session_state.get(f"editing_page_{idx}", False):
                                st.markdown(f"**Editando:** {pagina['titulo']}")
    
                                nuevo_titulo = st.text_input(
                                    "Título",
                                    value=pagina['titulo'],
                                    key=f"titulo_{idx}"
                                )
    
                                nuevo_contenido = st.text_area(
                                    "Contenido (Markdown)",
                                    value=pagina['contenido_markdown'],
                                    height=200,
                                    key=f"contenido_{idx}"
                                )
    
                                col_save, col_improve, col_cancel = st.columns([1, 1, 1])
    
                                with col_save:
                                    if st.button("💾 Guardar", key=f"save_{idx}", use_container_width=True):
                                        estructura['paginas'][idx]['titulo'] = nuevo_titulo
                                        estructura['paginas'][idx]['contenido_markdown'] = nuevo_contenido
                                        st.session_state.wiki_create_estructura_editada = estructura
                                        st.session_state[f"editing_page_{idx}"] = False
                                        st.success("✅ Guardado")
                                        st.rerun()
    
                                with col_improve:
                                    if st.button("✨ Mejorar con Frida", key=f"improve_{idx}", use_container_width=True):
                                        contenido_mejorado = mejorar_contenido_pagina_con_frida(
                                            nuevo_titulo,
                                            nuevo_contenido,
                                            st.session_state.wiki_create_doc_content[:3000]
                                        )
                                        estructura['paginas'][idx]['contenido_markdown'] = contenido_mejorado
                                        st.session_state.wiki_create_estructura_editada = estructura
                                        st.success("✅ Mejorado")
                                        st.rerun()
    
                                with col_cancel:
                                    if st.button("❌ Cancelar", key=f"cancel_{idx}", use_container_width=True):
                                        st.session_state[f"editing_page_{idx}"] = False
                                        st.rerun()
    
                            st.markdown("---")
    
                    st.markdown("---")
    
                    col_confirm, col_preview = st.columns([1, 1])
    
                    with col_confirm:
                        if st.button("✅ Confirmar Estructura", use_container_width=True, key="confirmar_estructura"):
                            st.session_state.wiki_create_ready_to_create = True
                            st.success("✅ Estructura confirmada. Pasa al siguiente paso.")
                            st.rerun()
    
                    with col_preview:
                        st.info(f"📊 Total: {len(estructura['paginas'])} páginas a crear")
    
            # === PASO 4: SELECCIONAR WIKI Y MODO DE CREACIÓN ===
            if st.session_state.wiki_create_ready_to_create:
                with st.expander("🎯 Paso 4: Configurar Creación en Azure DevOps", expanded=True):
                    st.markdown("**Selecciona la Wiki de destino y el modo de creación**")
    
                    # Listar wikis
                    col_wiki1, col_wiki2 = st.columns([2, 1])
    
                    with col_wiki1:
                        if st.button("🔍 Listar Wikis del Proyecto", key="listar_wikis_crear"):
                            with st.spinner("Obteniendo wikis..."):
                                wikis = obtener_wikis_proyecto(
                                    st.session_state.devops_org,
                                    st.session_state.devops_project,
                                    st.session_state.devops_pat
                                )
    
                            if wikis:
                                st.session_state.available_wikis_crear = wikis
                                st.success(f"✅ {len(wikis)} wiki(s) encontrada(s)")
                            else:
                                st.error("❌ No se encontraron wikis")
    
                        # Selector de wiki
                        if 'available_wikis_crear' in st.session_state and st.session_state.available_wikis_crear:
                            selected_wiki_idx = st.selectbox(
                                "Selecciona Wiki de destino",
                                options=range(len(st.session_state.available_wikis_crear)),
                                format_func=lambda i: f"{st.session_state.available_wikis_crear[i]['name']} ({st.session_state.available_wikis_crear[i]['type']})",
                                key="wiki_destino_selector"
                            )
    
                            selected_wiki = st.session_state.available_wikis_crear[selected_wiki_idx]
                            st.session_state.selected_wiki_id_crear = selected_wiki['id']
                            st.session_state.selected_wiki_name_crear = selected_wiki['name']
    
                            st.info(f"📖 Wiki seleccionada: **{selected_wiki['name']}**")
    
                    with col_wiki2:
                        # Modo de creación
                        modo_creacion = st.radio(
                            "Modo de creación",
                            options=["nueva", "extender"],
                            format_func=lambda x: "📄 Nueva página raíz" if x == "nueva" else "📁 Extender página existente",
                            key="modo_creacion_radio"
                        )
    
                        st.session_state.wiki_create_modo = modo_creacion
    
                    # Si modo extender, mostrar páginas existentes
                    if modo_creacion == "extender" and 'selected_wiki_id_crear' in st.session_state:
                        st.markdown("#### Selecciona página padre:")
    
                        if st.button("📋 Listar Páginas Existentes", key="listar_paginas_existentes_crear"):
                            with st.spinner("Obteniendo páginas..."):
                                estructura_existente = obtener_estructura_paginas_wiki_existente(
                                    st.session_state.devops_org,
                                    st.session_state.devops_project,
                                    st.session_state.devops_pat,
                                    st.session_state.selected_wiki_id_crear
                                )
    
                            if estructura_existente:
                                st.session_state.wiki_estructura_existente = estructura_existente
                                st.success(f"✅ {len(estructura_existente)} páginas encontradas")
    
                        if 'wiki_estructura_existente' in st.session_state and st.session_state.wiki_estructura_existente:
                            opciones_padre = ["/ (Raíz)"] + [f"{'　' * p['nivel']}{p['nombre']}" for p in st.session_state.wiki_estructura_existente]
                            paths_padre = ["/"] + [p['path'] for p in st.session_state.wiki_estructura_existente]
    
                            idx_padre = st.selectbox(
                                "Página padre",
                                options=range(len(opciones_padre)),
                                format_func=lambda i: opciones_padre[i],
                                key="padre_selector"
                            )
    
                            st.session_state.wiki_create_pagina_padre = paths_padre[idx_padre]
                            st.info(f"📌 Las páginas se crearán bajo: **{paths_padre[idx_padre]}**")
    
            # === PASO 5: CREAR PÁGINAS ===
            if st.session_state.wiki_create_ready_to_create and 'selected_wiki_id_crear' in st.session_state:
                st.markdown("---")
                st.markdown("### 🚀 Crear Páginas en Azure DevOps")
    
                col_warning, col_create = st.columns([2, 1])
    
                with col_warning:
                    st.warning("""
                    **⚠️ Atención:**
                    - Se crearán páginas reales en Azure DevOps
                    - Si ya existen páginas con el mismo nombre, se intentarán actualizar
                    - Revisa bien la estructura antes de continuar
                    """)
    
                    st.info(f"""
                    **Resumen de creación:**
                    - Wiki destino: {st.session_state.get('selected_wiki_name_crear', 'No seleccionada')}
                    - Modo: {"Nueva página raíz" if st.session_state.wiki_create_modo == "nueva" else f"Extender bajo {st.session_state.get('wiki_create_pagina_padre', '/')}"}
                    - Total de páginas: {len(st.session_state.wiki_create_estructura_editada['paginas'])}
                    """)
    
                with col_create:
                    if st.button("🚀 Crear Páginas en Wiki", use_container_width=True, type="primary", key="crear_paginas_finales"):
                        estructura = st.session_state.wiki_create_estructura_editada
                        paginas = estructura['paginas']
    
                        # Ordenar páginas por orden y jerarquía
                        paginas_ordenadas = sorted(paginas, key=lambda x: x.get('orden', 0))
    
                        progress_bar = st.progress(0)
                        status_text = st.empty()
    
                        exitos = 0
                        errores = 0
    
                        # Mapa para tracking de paths reales de páginas creadas
                        titulo_a_path = {}
    
                        for idx, pagina in enumerate(paginas_ordenadas):
                            progress_bar.progress((idx + 1) / len(paginas_ordenadas))
                            status_text.text(f"Creando: {pagina['titulo']} ({idx + 1}/{len(paginas_ordenadas)})")
    
                            titulo_clean = pagina['titulo'].replace(' ', '-')
    
                            # Construir path correctamente usando el mapa de paths
                            if st.session_state.wiki_create_modo == "nueva":
                                # Modo nueva: crear desde raíz
                                if pagina.get('es_raiz', False):
                                    path = f"/{titulo_clean}"
                                else:
                                    padre_titulo = pagina.get('padre', '')
                                    if padre_titulo and padre_titulo in titulo_a_path:
                                        # Usar el path real del padre desde el mapa
                                        path_padre = titulo_a_path[padre_titulo]
                                        path = f"{path_padre}/{titulo_clean}"
                                    else:
                                        # Fallback: asumir que está en la raíz
                                        path = f"/{titulo_clean}"
                            else:
                                # Modo extender: añadir bajo página padre
                                base_path = st.session_state.wiki_create_pagina_padre
                                if base_path == "/":
                                    path = f"/{titulo_clean}"
                                else:
                                    if pagina.get('es_raiz', False):
                                        path = f"{base_path}/{titulo_clean}"
                                    else:
                                        padre_titulo = pagina.get('padre', '')
                                        if padre_titulo and padre_titulo in titulo_a_path:
                                            # Usar el path real del padre
                                            path_padre = titulo_a_path[padre_titulo]
                                            path = f"{path_padre}/{titulo_clean}"
                                        else:
                                            path = f"{base_path}/{titulo_clean}"
    
                            # Crear la página
                            success, result = crear_pagina_wiki_azure(
                                st.session_state.devops_org,
                                st.session_state.devops_project,
                                st.session_state.devops_pat,
                                st.session_state.selected_wiki_id_crear,
                                path,
                                pagina['contenido_markdown']
                            )
    
                            if success:
                                exitos += 1
                                # Guardar el path real en el mapa
                                titulo_a_path[pagina['titulo']] = path
                                st.success(f"✅ Creada: {pagina['titulo']} → {path}")
                            else:
                                errores += 1
                                st.error(f"❌ Error: {pagina['titulo']}")
    
                        progress_bar.progress(1.0)
                        status_text.text("¡Creación completada!")
    
                        st.markdown("---")
                        st.success(f"""
                        **✅ Proceso completado**
                        - Páginas creadas: {exitos}
                        - Errores: {errores}
                        - Total procesadas: {len(paginas_ordenadas)}
                        """)
    
                        # Botón para limpiar y empezar de nuevo
                        if st.button("🔄 Crear otra wiki desde documento", key="reset_all_wiki_create"):
                            st.session_state.wiki_create_doc_content = ""
                            st.session_state.wiki_create_doc_filename = ""
                            st.session_state.wiki_create_estructura_propuesta = None
                            st.session_state.wiki_create_estructura_editada = None
                            st.session_state.wiki_create_ready_to_create = False
                            if 'available_wikis_crear' in st.session_state:
                                del st.session_state.available_wikis_crear
                            if 'wiki_estructura_existente' in st.session_state:
                                del st.session_state.wiki_estructura_existente
                            st.rerun()

    # ================= SUBTAB 4: CREAR TAREA =================
    with subtab_crear_tarea:
        st.subheader("➕ Crear Tarea en Azure DevOps")
        st.markdown("Crea work items automáticamente usando IA o manualmente")

        # Verificar configuración de Azure DevOps
        if not st.session_state.devops_pat or not st.session_state.devops_org or not st.session_state.devops_project:
            st.warning("⚠️ Primero configura la conexión a Azure DevOps en la sección de Configuración arriba")
        else:
            # Inicializar variables de sesión para esta funcionalidad
            if 'workitem_generated' not in st.session_state:
                st.session_state.workitem_generated = False
            if 'workitem_data' not in st.session_state:
                st.session_state.workitem_data = {}
            if 'custom_prompt_workitem' not in st.session_state:
                st.session_state.custom_prompt_workitem = ""
            if 'current_field_mappings' not in st.session_state:
                st.session_state.current_field_mappings = {}
            if 'previous_workitem_type' not in st.session_state:
                st.session_state.previous_workitem_type = "User Story"

            # === PASO 1: GENERAR CON IA (OPCIONAL) ===
            with st.expander("🤖 Paso 1 (Opcional): Generar campos con IA", expanded=not st.session_state.workitem_generated):
                st.markdown("**Describe la tarea que quieres crear y la IA completará todos los campos automáticamente**")

                col_plantilla, col_desc = st.columns([1, 2])

                with col_plantilla:
                    template_choice = st.selectbox(
                        "Plantilla de contexto",
                        options=[
                            "PO Crear Work Item",
                            "PO Definicion historia",
                            "PO Definicion epica",
                            "PO Definicion mejora tecnica",
                            "PO Definicion spike",
                            "Libre"
                        ],
                        help="Selecciona una plantilla para guiar a la IA",
                        key="workitem_template_choice"
                    )

                with col_desc:
                    descripcion_ia = st.text_area(
                        "Descripción de la tarea",
                        height=150,
                        placeholder="Ejemplo: Crear una funcionalidad para exportar reportes en PDF...",
                        help="Describe lo que quieres que haga la tarea",
                        key="workitem_descripcion_ia"
                    )

                # Área de texto editable para el prompt
                st.markdown("---")
                st.markdown("**✏️ Prompt que se enviará a la IA (editable)**")

                # Generar el prompt base basado en la plantilla seleccionada
                template = get_template(template_choice)
                if descripcion_ia:
                    prompt_preview = template.format(input=descripcion_ia)
                else:
                    prompt_preview = template.replace("{input}", "[Tu descripción aquí]")

                # Adaptar instrucción del campo descripción según la plantilla seleccionada
                if template_choice == "PO Definicion historia":
                    desc_field_instructions = (
                        "Descripción detallada en formato HTML con los puntos principales. "
                        "Quiero en este mismo campo primero detallar la historia con forma: "
                        "Como; Quiero; Para. "
                        "Quiero a continuacion tambien una descripcion general de la historia. "
                        "Tambien han de mostrarse los casos de uso con el formato: dado, cuando y entonces. "
                        "Necesito que los casos de uso, se muestren en una tabla de 3 columnas."
                    )
                elif template_choice == "PO Definicion epica":
                    desc_field_instructions = (
                        "Descripción detallada en formato HTML con los puntos principales. "
                        "Quiero detallar una epica con forma: Creemos que; Para; Conseguiremos. "
                        "A continuacion Quiero tambien una descripcion de la epica."
                    )
                else:
                    desc_field_instructions = (
                        "Descripción detallada en formato HTML. Incluye TODA la información generada anteriormente "
                        "(tablas, casos de uso, etc.) convertida a HTML usando "
                        "div, p, ul, li, ol, strong, table, tr, td, etc."
                    )

                # Agregar instrucciones JSON al final
                json_instructions = ("""

IMPORTANTE: Además de todo lo anterior, debes devolver AL FINAL de tu respuesta un bloque JSON válido con la siguiente estructura EXACTA:

```json
{
    "titulo": "Título conciso y claro del work item",
    "descripcion": "DESCRIPCION_PLACEHOLDER",
    "acceptance_criteria": "Criterios de aceptación en formato HTML. Si generaste una tabla de criterios, conviértela a HTML.",
    "dependencies": "Dependencias identificadas en formato HTML. Si generaste una tabla de dependencias, conviértela a HTML. Si no hay, cadena vacía.",
    "riesgos": "Riesgos potenciales en formato HTML. Si generaste una tabla de riesgos, conviértela a HTML. Si no hay, cadena vacía.",
    "team": "Equipo responsable sugerido (si aplica, sino vacío)",
    "source": "Origen de la tarea (si aplica, sino vacío)",
    "value_area": "Una de estas opciones: Business, Architectural, Design, Development"
}
```

REGLAS CRÍTICAS:
1. El JSON debe estar AL FINAL de tu respuesta completa
2. Toda la información detallada (tablas, criterios, casos de uso, etc.) debe ir en los campos HTML del JSON
3. Usa formato HTML válido en descripcion, acceptance_criteria, dependencies y riesgos
4. El JSON debe ser parseable y válido
5. Si no hay información para un campo opcional, usa cadena vacía ""
6. IMPORTANTE: dentro del contenido HTML NO uses comillas dobles ("). Usa &quot; para comillas en el texto y comillas simples (') para atributos HTML
""").replace("DESCRIPCION_PLACEHOLDER", desc_field_instructions)

                full_prompt = prompt_preview + json_instructions

                # Detectar cambios en descripcion o plantilla para actualizar el text_area.
                # Sin esto, Streamlit ignora el valor recalculado de full_prompt porque
                # la clave ya existe en session_state, y la IA recibe el prompt sin el
                # texto del usuario.
                prompt_key = f"workitem_custom_prompt_{template_choice}"
                prev_desc_key = "_prev_workitem_desc"
                prev_tmpl_key = "_prev_workitem_tmpl"
                if (st.session_state.get(prev_desc_key) != descripcion_ia or
                        st.session_state.get(prev_tmpl_key) != template_choice):
                    st.session_state[prompt_key] = full_prompt
                    st.session_state[prev_desc_key] = descripcion_ia
                    st.session_state[prev_tmpl_key] = template_choice

                custom_prompt = st.text_area(
                    "Prompt completo",
                    height=300,
                    help="Puedes editar este prompt antes de enviarlo a la IA",
                    key=prompt_key
                )

                # Mostrar información del modelo antes del botón
                if 'model' in st.session_state:
                    st.info(f"📊 Modelo configurado: **{st.session_state.model}**")
                else:
                    st.error("❌ No hay modelo configurado. Por favor configura el modelo en el sidebar.")

                if st.button("🎯 Generar campos con IA", disabled=not descripcion_ia, key="workitem_generar_ia_btn"):
                    st.write("🔄 **Botón presionado - iniciando proceso...**")
                    with st.spinner("🧠 Frida está generando los campos de la tarea..."):
                        try:
                            # Usar el prompt del text_area (siempre actualizado)
                            prompt_to_use = custom_prompt

                            # Verificar que el modelo esté configurado
                            if 'model' not in st.session_state:
                                st.error("❌ No hay modelo configurado. Por favor configura el modelo en el sidebar.")
                                st.stop()

                            payload = {
                                "model": st.session_state.model,
                                "messages": [
                                    {"role": "system", "content": "Eres un experto Product Owner. Genera una respuesta detallada según la plantilla, y AL FINAL incluye un JSON válido con los campos estructurados."},
                                    {"role": "user", "content": prompt_to_use}
                                ]
                            }

                            # Debug: mostrar payload
                            st.info(f"🔍 Enviando petición a la IA con modelo: {st.session_state.model}")

                            response = call_ia(payload)

                            # Debug: mostrar que se recibió respuesta
                            st.info(f"✅ Respuesta recibida ({len(response)} caracteres)")

                            # Extraer el JSON de la respuesta
                            response_text = response.strip()
                            json_str = None

                            # Intentar encontrar JSON entre ```json y ```
                            if "```json" in response_text:
                                json_start = response_text.rfind("```json") + 7
                                json_end = response_text.find("```", json_start)
                                if json_end != -1:
                                    json_str = response_text[json_start:json_end].strip()
                                    st.info("🔍 JSON encontrado en formato markdown")

                            # Si no se encontró con markdown, buscar entre llaves
                            if not json_str:
                                last_open = response_text.rfind("{")
                                if last_open != -1:
                                    brace_count = 0
                                    for i in range(last_open, len(response_text)):
                                        if response_text[i] == "{":
                                            brace_count += 1
                                        elif response_text[i] == "}":
                                            brace_count -= 1
                                            if brace_count == 0:
                                                json_str = response_text[last_open:i+1]
                                                st.info("🔍 JSON encontrado entre llaves")
                                                break

                            if not json_str:
                                st.error("❌ No se encontró un bloque JSON válido en la respuesta")
                                with st.expander("🔍 Ver respuesta completa de la IA"):
                                    st.code(response_text)
                                raise ValueError("No se encontró un bloque JSON válido en la respuesta")

                            # Parsear el JSON (con fallback para comillas sin escapar en HTML)
                            try:
                                data = json.loads(json_str)
                            except json.JSONDecodeError:
                                sanitized = sanitize_json_string(json_str)
                                data = json.loads(sanitized)
                            st.info("✅ JSON parseado correctamente")

                            # Guardar en session_state
                            st.session_state.workitem_data = {
                                'titulo': data.get('titulo', ''),
                                'descripcion': data.get('descripcion', ''),
                                'acceptance_criteria': data.get('acceptance_criteria', ''),
                                'dependencies': data.get('dependencies', ''),
                                'riesgos': data.get('riesgos', ''),
                                'team': data.get('team', ''),
                                'source': data.get('source', ''),
                                'value_area': data.get('value_area', 'Business')
                            }

                            # Actualizar también las claves de los widgets directamente,
                            # porque Streamlit ignora el parámetro `value=` si la clave
                            # ya existe en session_state (que es el caso en reruns).
                            for campo in ['titulo', 'descripcion', 'acceptance_criteria',
                                          'dependencies', 'riesgos', 'team', 'source', 'value_area']:
                                st.session_state[f'value_{campo}'] = st.session_state.workitem_data[campo]
                            st.session_state.workitem_generated = True
                            st.success("✅ Campos generados correctamente. Revísalos abajo y modifícalos si es necesario.")
                            st.rerun()

                        except json.JSONDecodeError as e:
                            st.error(f"❌ Error al parsear el JSON: {str(e)}")
                            if 'response' in locals():
                                with st.expander("Ver respuesta de la IA"):
                                    st.code(response)
                            st.info("💡 Intenta usar una descripción más específica o prueba con otra plantilla")
                        except ValueError as e:
                            st.error(f"❌ {str(e)}")
                            if 'response' in locals():
                                with st.expander("Ver respuesta de la IA"):
                                    st.code(response)
                            st.info("💡 La IA no devolvió un JSON válido. Intenta con una descripción más clara")
                        except Exception as e:
                            st.error(f"❌ Error inesperado: {type(e).__name__}: {str(e)}")
                            import traceback
                            st.code(traceback.format_exc())
                            if 'response' in locals():
                                with st.expander("Ver respuesta de la IA"):
                                    st.code(response)

            # Mostrar respuesta JSON de la IA (desplegable discreto)
            if st.session_state.workitem_generated and st.session_state.workitem_data:
                with st.expander("🔍 Ver respuesta JSON de la IA", expanded=False):
                    st.code(json.dumps(st.session_state.workitem_data, indent=2, ensure_ascii=False), language="json")

            # === PASO 2: COMPLETAR/EDITAR CAMPOS ===
            st.markdown("---")
            st.markdown("### 📝 Paso 2: Completar o editar campos")

            # Columnas para organizar el formulario
            col_tipo, col_area, col_iteration = st.columns(3)

            with col_tipo:
                work_item_type = st.selectbox(
                    "Tipo de tarea *",
                    options=["User Story", "Feature", "Epic", "Bug", "Task"],
                    help="Tipo de work item a crear",
                    key="workitem_type_selector"
                )

            # Detectar cambio de tipo de tarea y actualizar mappings
            if work_item_type != st.session_state.previous_workitem_type:
                st.session_state.previous_workitem_type = work_item_type

                import copy
                new_mappings = copy.deepcopy(
                    WORKITEM_FIELD_MAPPING.get(work_item_type, WORKITEM_FIELD_MAPPING["User Story"])
                )
                st.session_state.current_field_mappings = new_mappings

                # Actualizar las claves de los widgets de Azure field y enabled,
                # porque Streamlit ignora value= si la clave ya existe en session_state.
                for fname, fmap in new_mappings.items():
                    st.session_state[f"azure_{fname}"] = fmap.get('azure_field', '')
                    st.session_state[f"enable_{fname}"] = fmap.get('enabled', False)

                st.rerun()

            # Cargar field mappings si están vacíos (primera vez)
            if not st.session_state.current_field_mappings:
                import copy
                new_mappings = copy.deepcopy(
                    WORKITEM_FIELD_MAPPING.get(work_item_type, WORKITEM_FIELD_MAPPING["User Story"])
                )
                st.session_state.current_field_mappings = new_mappings
                for fname, fmap in new_mappings.items():
                    st.session_state[f"azure_{fname}"] = fmap.get('azure_field', '')
                    st.session_state[f"enable_{fname}"] = fmap.get('enabled', False)

            with col_area:
                area_path = st.text_input(
                    "Área (Area Path)",
                    value="",
                    placeholder="Ejemplo: Sales\\MySaga POC",
                    help="Área del proyecto. Deja vacío para usar la raíz del proyecto",
                    key="workitem_area_path"
                )

            with col_iteration:
                iteration_path = st.text_input(
                    "Iteración (Iteration Path)",
                    value="",
                    placeholder="Ejemplo: Sprint 1",
                    help="Iteración/Sprint. Deja vacío para la iteración por defecto",
                    key="workitem_iteration_path"
                )

            st.markdown("---")
            st.markdown("#### 🔧 Configuración de campos y mapeo a Azure DevOps")
            st.markdown("*Activa/desactiva campos y edita los nombres de campos en Azure DevOps*")

            # Diccionario para almacenar los estados actualizados de los campos
            updated_mappings = {}

            # Definir labels amigables para cada campo
            field_labels = {
                'titulo': 'Título',
                'descripcion': 'Descripción',
                'acceptance_criteria': 'Criterios de Aceptación',
                'dependencies': 'Dependencias',
                'riesgos': 'Riesgos',
                'team': 'Team',
                'source': 'Source',
                'value_area': 'Value Area'
            }

            # Campos con sus valores
            field_values = {}

            # TITULO (siempre requerido)
            st.markdown("##### Campos Obligatorios")

            # Título
            field_name = 'titulo'
            mapping = st.session_state.current_field_mappings.get(field_name, {"azure_field": "System.Title", "enabled": True})

            col_check, col_content = st.columns([0.5, 10])
            with col_check:
                titulo_enabled = st.checkbox(
                    "Habilitar",
                    value=mapping.get('enabled', True),
                    key=f"enable_{field_name}",
                    help="Este campo es obligatorio",
                    disabled=True,  # Título siempre habilitado
                    label_visibility="collapsed"
                )
            with col_content:
                col_label, col_azure, col_value = st.columns([2, 3, 5])
                with col_label:
                    st.markdown(f"**{field_labels[field_name]}** *")
                with col_azure:
                    titulo_azure_field = st.text_input(
                        "Campo Azure",
                        value=mapping.get('azure_field', 'System.Title'),
                        key=f"azure_{field_name}",
                        label_visibility="collapsed",
                        placeholder="System.Title"
                    )
                with col_value:
                    titulo_value = st.text_input(
                        "Valor",
                        value=st.session_state.workitem_data.get(field_name, ''),
                        placeholder="Título conciso de la tarea",
                        key=f"value_{field_name}",
                        label_visibility="collapsed"
                    )

            updated_mappings[field_name] = {
                "azure_field": titulo_azure_field,
                "enabled": titulo_enabled
            }
            field_values[field_name] = titulo_value

            # Descripción
            field_name = 'descripcion'
            mapping = st.session_state.current_field_mappings.get(field_name, {"azure_field": "System.Description", "enabled": True})

            col_check, col_content = st.columns([0.5, 10])
            with col_check:
                descripcion_enabled = st.checkbox(
                    "Habilitar",
                    value=mapping.get('enabled', True),
                    key=f"enable_{field_name}",
                    help="Este campo es obligatorio",
                    disabled=True,  # Descripción siempre habilitada
                    label_visibility="collapsed"
                )
            with col_content:
                col_label, col_azure = st.columns([2, 8])
                with col_label:
                    st.markdown(f"**{field_labels[field_name]}** *")
                with col_azure:
                    descripcion_azure_field = st.text_input(
                        "Campo Azure",
                        value=mapping.get('azure_field', 'System.Description'),
                        key=f"azure_{field_name}",
                        label_visibility="collapsed",
                        placeholder="System.Description"
                    )
                descripcion_value = st.text_area(
                    "Valor",
                    value=st.session_state.workitem_data.get(field_name, ''),
                    height=150,
                    placeholder="Descripción detallada de la tarea (puede usar HTML)...",
                    key=f"value_{field_name}",
                    label_visibility="collapsed"
                )

            updated_mappings[field_name] = {
                "azure_field": descripcion_azure_field,
                "enabled": descripcion_enabled
            }
            field_values[field_name] = descripcion_value

            # Campos opcionales
            st.markdown("##### Campos Opcionales")

            # Acceptance Criteria
            field_name = 'acceptance_criteria'
            mapping = st.session_state.current_field_mappings.get(field_name, {"azure_field": "", "enabled": False})

            col_check, col_content = st.columns([0.5, 10])
            with col_check:
                ac_enabled = st.checkbox(
                    "Habilitar",
                    value=mapping.get('enabled', False),
                    key=f"enable_{field_name}",
                    label_visibility="collapsed"
                )
            with col_content:
                col_label, col_azure = st.columns([2, 8])
                with col_label:
                    st.markdown(f"**{field_labels[field_name]}**")
                with col_azure:
                    ac_azure_field = st.text_input(
                        "Campo Azure",
                        value=mapping.get('azure_field', 'Microsoft.VSTS.Common.AcceptanceCriteria'),
                        key=f"azure_{field_name}",
                        label_visibility="collapsed",
                        placeholder="Microsoft.VSTS.Common.AcceptanceCriteria",
                        disabled=not ac_enabled
                    )
                ac_value = st.text_area(
                    "Valor",
                    value=st.session_state.workitem_data.get(field_name, ''),
                    height=120,
                    placeholder="Criterios de aceptación (puede usar HTML)...",
                    key=f"value_{field_name}",
                    label_visibility="collapsed",
                    disabled=not ac_enabled
                )

            updated_mappings[field_name] = {
                "azure_field": ac_azure_field,
                "enabled": ac_enabled
            }
            field_values[field_name] = ac_value if ac_enabled else ""

            # Dependencies
            field_name = 'dependencies'
            mapping = st.session_state.current_field_mappings.get(field_name, {"azure_field": "", "enabled": False})

            col_check, col_content = st.columns([0.5, 10])
            with col_check:
                dep_enabled = st.checkbox(
                    "Habilitar",
                    value=mapping.get('enabled', False),
                    key=f"enable_{field_name}",
                    label_visibility="collapsed"
                )
            with col_content:
                col_label, col_azure = st.columns([2, 8])
                with col_label:
                    st.markdown(f"**{field_labels[field_name]}**")
                with col_azure:
                    dep_azure_field = st.text_input(
                        "Campo Azure",
                        value=mapping.get('azure_field', 'Custom.Dependencies'),
                        key=f"azure_{field_name}",
                        label_visibility="collapsed",
                        placeholder="Custom.Dependencies",
                        disabled=not dep_enabled
                    )
                dep_value = st.text_area(
                    "Valor",
                    value=st.session_state.workitem_data.get(field_name, ''),
                    height=100,
                    placeholder="Dependencias con otras tareas o sistemas...",
                    key=f"value_{field_name}",
                    label_visibility="collapsed",
                    disabled=not dep_enabled
                )

            updated_mappings[field_name] = {
                "azure_field": dep_azure_field,
                "enabled": dep_enabled
            }
            field_values[field_name] = dep_value if dep_enabled else ""

            # Riesgos
            field_name = 'riesgos'
            mapping = st.session_state.current_field_mappings.get(field_name, {"azure_field": "", "enabled": False})

            col_check, col_content = st.columns([0.5, 10])
            with col_check:
                riesgos_enabled = st.checkbox(
                    "Habilitar",
                    value=mapping.get('enabled', False),
                    key=f"enable_{field_name}",
                    label_visibility="collapsed"
                )
            with col_content:
                col_label, col_azure = st.columns([2, 8])
                with col_label:
                    st.markdown(f"**{field_labels[field_name]}**")
                with col_azure:
                    riesgos_azure_field = st.text_input(
                        "Campo Azure",
                        value=mapping.get('azure_field', 'Custom.Riesgos'),
                        key=f"azure_{field_name}",
                        label_visibility="collapsed",
                        placeholder="Custom.Riesgos",
                        disabled=not riesgos_enabled
                    )
                riesgos_value = st.text_area(
                    "Valor",
                    value=st.session_state.workitem_data.get(field_name, ''),
                    height=100,
                    placeholder="Riesgos identificados...",
                    key=f"value_{field_name}",
                    label_visibility="collapsed",
                    disabled=not riesgos_enabled
                )

            updated_mappings[field_name] = {
                "azure_field": riesgos_azure_field,
                "enabled": riesgos_enabled
            }
            field_values[field_name] = riesgos_value if riesgos_enabled else ""

            # Team
            field_name = 'team'
            mapping = st.session_state.current_field_mappings.get(field_name, {"azure_field": "", "enabled": False})

            col_check, col_content = st.columns([0.5, 10])
            with col_check:
                team_enabled = st.checkbox(
                    "Habilitar",
                    value=mapping.get('enabled', False),
                    key=f"enable_{field_name}",
                    label_visibility="collapsed"
                )
            with col_content:
                col_label, col_azure, col_value = st.columns([2, 3, 5])
                with col_label:
                    st.markdown(f"**{field_labels[field_name]}**")
                with col_azure:
                    team_azure_field = st.text_input(
                        "Campo Azure",
                        value=mapping.get('azure_field', 'Custom.Team'),
                        key=f"azure_{field_name}",
                        label_visibility="collapsed",
                        placeholder="Custom.Team",
                        disabled=not team_enabled
                    )
                with col_value:
                    team_value = st.text_input(
                        "Valor",
                        value=st.session_state.workitem_data.get(field_name, ''),
                        placeholder="Equipo responsable",
                        key=f"value_{field_name}",
                        label_visibility="collapsed",
                        disabled=not team_enabled
                    )

            updated_mappings[field_name] = {
                "azure_field": team_azure_field,
                "enabled": team_enabled
            }
            field_values[field_name] = team_value if team_enabled else ""

            # Source
            field_name = 'source'
            mapping = st.session_state.current_field_mappings.get(field_name, {"azure_field": "", "enabled": False})

            col_check, col_content = st.columns([0.5, 10])
            with col_check:
                source_enabled = st.checkbox(
                    "Habilitar",
                    value=mapping.get('enabled', False),
                    key=f"enable_{field_name}",
                    label_visibility="collapsed"
                )
            with col_content:
                col_label, col_azure, col_value = st.columns([2, 3, 5])
                with col_label:
                    st.markdown(f"**{field_labels[field_name]}**")
                with col_azure:
                    source_azure_field = st.text_input(
                        "Campo Azure",
                        value=mapping.get('azure_field', 'Custom.Source'),
                        key=f"azure_{field_name}",
                        label_visibility="collapsed",
                        placeholder="Custom.Source",
                        disabled=not source_enabled
                    )
                with col_value:
                    source_value = st.text_input(
                        "Valor",
                        value=st.session_state.workitem_data.get(field_name, ''),
                        placeholder="Origen de la tarea",
                        key=f"value_{field_name}",
                        label_visibility="collapsed",
                        disabled=not source_enabled
                    )

            updated_mappings[field_name] = {
                "azure_field": source_azure_field,
                "enabled": source_enabled
            }
            field_values[field_name] = source_value if source_enabled else ""

            # Value Area
            field_name = 'value_area'
            mapping = st.session_state.current_field_mappings.get(field_name, {"azure_field": "", "enabled": False})

            col_check, col_content = st.columns([0.5, 10])
            with col_check:
                va_enabled = st.checkbox(
                    "Habilitar",
                    value=mapping.get('enabled', False),
                    key=f"enable_{field_name}",
                    label_visibility="collapsed"
                )
            with col_content:
                col_label, col_azure, col_value = st.columns([2, 3, 5])
                with col_label:
                    st.markdown(f"**{field_labels[field_name]}**")
                with col_azure:
                    va_azure_field = st.text_input(
                        "Campo Azure",
                        value=mapping.get('azure_field', 'Microsoft.VSTS.Common.ValueArea'),
                        key=f"azure_{field_name}",
                        label_visibility="collapsed",
                        placeholder="Microsoft.VSTS.Common.ValueArea",
                        disabled=not va_enabled
                    )
                with col_value:
                    current_value = st.session_state.workitem_data.get(field_name, 'Business')
                    try:
                        current_index = ["Business", "Architectural", "Design", "Development"].index(current_value)
                    except ValueError:
                        current_index = 0

                    va_value = st.selectbox(
                        "Valor",
                        options=["Business", "Architectural", "Design", "Development"],
                        index=current_index,
                        key=f"value_{field_name}",
                        label_visibility="collapsed",
                        disabled=not va_enabled
                    )

            updated_mappings[field_name] = {
                "azure_field": va_azure_field,
                "enabled": va_enabled
            }
            field_values[field_name] = va_value if va_enabled else ""

            # Actualizar los mappings en session_state
            st.session_state.current_field_mappings = updated_mappings

            # === PASO 3: CREAR EN AZURE DEVOPS ===
            st.markdown("---")
            st.markdown("### ✅ Paso 3: Crear en Azure DevOps")

            col_btn, col_reset = st.columns([3, 1])

            with col_btn:
                if st.button("🚀 Crear Work Item en Azure DevOps", type="primary", use_container_width=True, key="workitem_crear_btn"):
                    # Validar campos obligatorios
                    if not field_values['titulo'] or not field_values['descripcion']:
                        st.error("❌ El Título y la Descripción son obligatorios")
                    else:
                        with st.spinner("📤 Creando work item en Azure DevOps..."):
                            # Preparar campos con todos los valores
                            campos = field_values.copy()

                            # Agregar area e iteration si se especificaron
                            if area_path:
                                campos['area_path'] = area_path
                            if iteration_path:
                                campos['iteration_path'] = iteration_path

                            # Crear el work item con field_mappings
                            result = crear_workitem_devops(
                                st.session_state.devops_org,
                                st.session_state.devops_project,
                                st.session_state.devops_pat,
                                work_item_type,
                                campos,
                                st.session_state.current_field_mappings
                            )

                            if result['success']:
                                st.success(f"✅ **Work Item creado exitosamente!**")
                                st.info(f"**ID:** {result['id']}")
                                st.markdown(f"**URL:** [{result['url']}]({result['url']})")
                                st.balloons()

                                # Limpiar datos
                                st.session_state.workitem_generated = False
                                st.session_state.workitem_data = {}
                                st.session_state.custom_prompt_workitem = ""
                                st.session_state.current_field_mappings = {}
                            else:
                                st.error(f"❌ Error al crear work item: {result['error']}")

            with col_reset:
                if st.button("🔄 Limpiar", use_container_width=True, key="workitem_limpiar_btn"):
                    st.session_state.workitem_generated = False
                    st.session_state.workitem_data = {}
                    st.session_state.custom_prompt_workitem = ""
                    st.session_state.current_field_mappings = {}
                    st.rerun()

    # ================= TAB 3: ANÁLISIS DOCUMENTOS =================
with tab_doc:
    st.title("📄 Análisis de Documentos")
    st.markdown("Carga un documento Word y haz preguntas sobre su contenido o genera work items automáticamente")
    
    # Configuración del documento
    with st.expander("📥 Cargar Documento", expanded=not st.session_state.doc_indexed):
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("#### Opción 1: Subir archivo local")
            uploaded_doc = st.file_uploader(
                "Sube un documento Word (.docx)", 
                type=["docx"],
                help="Archivo .docx desde tu ordenador",
                key="upload_doc_file"
            )
            
            st.markdown("#### Opción 2: Desde Azure DevOps Work Item")
            col_az1, col_az2 = st.columns(2)
            
            with col_az1:
                workitem_id_input = st.number_input(
                    "ID de Work Item",
                    min_value=0,
                    value=0,
                    step=1,
                    help="ID del work item que contiene el documento adjunto"
                )
            
            with col_az2:
                if st.button("📋 Ver documentos adjuntos", disabled=(workitem_id_input == 0)):
                    if not st.session_state.devops_pat:
                        st.error("❌ Primero configura Azure DevOps en la pestaña 'Consulta Tareas'")
                    else:
                        with st.spinner("Obteniendo adjuntos..."):
                            attachments = obtener_attachments_workitem(
                                st.session_state.devops_org or "TelepizzaIT",
                                st.session_state.devops_project or "Sales",
                                st.session_state.devops_pat,
                                int(workitem_id_input)
                            )
                        
                        if attachments:
                            st.session_state.temp_attachments = attachments
                            st.success(f"✅ {len(attachments)} documento(s) .docx encontrado(s)")
                        else:
                            st.warning("⚠️ No se encontraron documentos .docx en este work item")
            
            # Mostrar lista de attachments si existen
            if st.session_state.temp_attachments:
                st.markdown("**Documentos Word encontrados:**")
                selected_attachment = st.selectbox(
                    "Selecciona un documento",
                    options=range(len(st.session_state.temp_attachments)),
                    format_func=lambda i: st.session_state.temp_attachments[i]["name"],
                    key="selected_attachment_idx"
                )
                
                st.session_state.selected_attachment_url = st.session_state.temp_attachments[selected_attachment]["url"]
                st.session_state.selected_attachment_name = st.session_state.temp_attachments[selected_attachment]["name"]
            
            st.markdown("#### Opción 3: URL pública")
            doc_url = st.text_input(
                "URL del documento",
                placeholder="https://ejemplo.com/documento.docx",
                help="URL de acceso público a un archivo .docx"
            )
        
        with col2:
            st.markdown("#### ⚙️ Configuración")
            chunk_size = st.slider(
                "Tamaño de fragmentos",
                min_value=500,
                max_value=4000,
                value=1000,
                step=100,
                help="Tamaño de cada fragmento del documento"
            )
            
            doc_top_k = st.slider(
                "Fragmentos relevantes",
                min_value=2,
                max_value=5,
                value=3,
                step=1,
                help="Número de fragmentos a usar como contexto"
            )
        
        st.markdown("---")
        
        col_btn1, col_btn2 = st.columns([3, 1])
        
        with col_btn1:
            if st.button("🔄 Procesar Documento", use_container_width=True, key="procesar_doc_btn"):
                doc_bytes = None
                filename = ""
                
                # Opción 1: Archivo local
                if uploaded_doc is not None:
                    doc_bytes = uploaded_doc.read()
                    filename = uploaded_doc.name
                    st.info(f"📄 Procesando: {filename}")
                
                # Opción 2: Desde Azure DevOps
                elif st.session_state.selected_attachment_url:
                    if not st.session_state.devops_pat:
                        st.error("❌ Configura Azure DevOps primero")
                    else:
                        with st.spinner("📥 Descargando desde Azure DevOps..."):
                            doc_bytes = descargar_attachment_devops(
                                st.session_state.selected_attachment_url,
                                st.session_state.devops_pat
                            )
                        filename = st.session_state.selected_attachment_name
                        if doc_bytes:
                            st.info(f"📄 Procesando: {filename}")
                
                # Opción 3: URL pública
                elif doc_url:
                    with st.spinner("📥 Descargando documento desde URL..."):
                        doc_bytes = descargar_documento_url(doc_url)
                    filename = doc_url.split("/")[-1]
                
                else:
                    st.error("❌ Debes subir un archivo, seleccionar uno de Azure DevOps o proporcionar una URL")
                
                if doc_bytes:
                    # Leer contenido
                    with st.spinner("📖 Leyendo contenido del documento..."):
                        contenido = leer_docx_desde_bytes(doc_bytes)
                    
                    if contenido:
                        st.success(f"✅ Documento leído: {len(contenido)} caracteres")
                        
                        # Dividir en chunks
                        chunks = dividir_en_chunks(contenido, chunk_size=chunk_size)
                        st.info(f"📑 Dividido en {len(chunks)} fragmentos")
                        
                        # Cargar modelo si no está cargado
                        if st.session_state.embedding_model is None:
                            st.session_state.embedding_model = cargar_modelo_embeddings()
                        
                        # Generar embeddings
                        embeddings = generar_embeddings_documento(
                            chunks,
                            st.session_state.embedding_model
                        )
                        
                        # Guardar en session_state
                        st.session_state.doc_content = contenido
                        st.session_state.doc_chunks = chunks
                        st.session_state.doc_embeddings = embeddings
                        st.session_state.doc_indexed = True
                        st.session_state.doc_filename = filename
                        st.session_state.doc_top_k = doc_top_k
                        
                        # Limpiar attachments temporales
                        st.session_state.temp_attachments = []
                        st.session_state.selected_attachment_url = ""
                        st.session_state.selected_attachment_name = ""
                        
                        st.success("✅ Documento indexado. Ya puedes hacer consultas o generar work items.")
                        st.rerun()
                    else:
                        st.error("❌ No se pudo leer el contenido del documento")
        
        with col_btn2:
            if st.button("🗑️ Limpiar", use_container_width=True, key="limpiar_doc"):
                st.session_state.doc_content = ""
                st.session_state.doc_chunks = []
                st.session_state.doc_embeddings = None
                st.session_state.doc_indexed = False
                st.session_state.doc_messages = []
                st.session_state.doc_filename = ""
                st.session_state.temp_attachments = []
                st.session_state.selected_attachment_url = ""
                st.session_state.selected_attachment_name = ""
                st.success("✅ Documento eliminado")
                st.rerun()
    
    # Estado del documento
    if st.session_state.doc_indexed:
        st.info(f"📄 **Documento cargado**: {st.session_state.doc_filename} ({len(st.session_state.doc_chunks)} fragmentos)")
        st.info(f"🎯 **Fragmentos por consulta**: {st.session_state.get('doc_top_k', 3)}")
    
    st.markdown("---")
    
    # Pestañas de funcionalidad
    subtab_chat, subtab_generate = st.tabs(["💬 Consultas", "🔧 Generar Work Items"])
    
    # SUBTAB: Consultas sobre el documento
    with subtab_chat:
        st.subheader("💬 Haz preguntas sobre el documento")
        
        for m in st.session_state.doc_messages:
            with st.chat_message(m["role"]):
                st.markdown(m["content"])
        
        if doc_query := st.chat_input(
            "Pregunta sobre el documento... ej: '¿Cuáles son los requisitos principales?'",
            key="doc_chat",
            disabled=not st.session_state.doc_indexed
        ):
            if not st.session_state.doc_indexed:
                st.warning("⚠️ Primero debes cargar y procesar un documento")
            else:
                st.session_state.doc_messages.append({"role": "user", "content": doc_query})
                
                top_k = st.session_state.get('doc_top_k', 3)
                
                # Buscar fragmentos relevantes
                with st.spinner(f"🔍 Buscando fragmentos relevantes..."):
                    resultados = buscar_chunks_similares(
                        doc_query,
                        st.session_state.doc_chunks,
                        st.session_state.doc_embeddings,
                        st.session_state.embedding_model,
                        top_k=top_k
                    )
                
                # Construir contexto
                contexto = construir_contexto_documento(resultados)
                
                # Llamar a Frida
                system_prompt = """Eres un asistente experto en analizar documentos técnicos y de negocio.

Cuando respondas:
1. Basa tu respuesta SOLO en la información de los fragmentos proporcionados
2. Si la información no está en los fragmentos, di que no está disponible en el documento
3. Sé preciso y cita partes específicas cuando sea relevante
4. Estructura tu respuesta de forma clara y concisa"""

                payload = {
                    "model": st.session_state.model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"{contexto}\n\n**Pregunta:** {doc_query}"}
                    ]
                }
                
                if st.session_state.include_temp:
                    payload["temperature"] = st.session_state.temperature
                if st.session_state.include_tokens:
                    payload["max_tokens"] = st.session_state.max_tokens
                
                with st.spinner("🤖 Frida está analizando el documento..."):
                    respuesta = call_ia(payload)
                
                st.session_state.doc_messages.append({"role": "assistant", "content": respuesta})
                
                # Mostrar fragmentos usados
                with st.expander(f"📋 Ver fragmentos utilizados"):
                    for i, resultado in enumerate(resultados, 1):
                        sim = resultado["similitud"]
                        chunk = resultado["chunk"]
                        
                        st.markdown(f"### Fragmento {i} (Relevancia: {sim:.1%})")
                        st.text(chunk[:500] + ("..." if len(chunk) > 500 else ""))
                        st.markdown("---")
                
                st.rerun()
    
    # SUBTAB: Generar work items
    with subtab_generate:
        st.subheader("🔧 Generar Work Items desde el Documento")
        
        if not st.session_state.doc_indexed:
            st.warning("⚠️ Primero debes cargar y procesar un documento")
        else:
            col_gen1, col_gen2 = st.columns([2, 1])
            
            with col_gen1:
                instruccion_generacion = st.text_area(
                    "Instrucciones de generación",
                    placeholder="Ej: Genera una épica principal con 3 historias de usuario basándote en los requisitos del documento",
                    height=100
                )
            
            with col_gen2:
                template_generacion = st.selectbox(
                    "Template a usar",
                    [
                        "PO Definicion epica",
                        "PO Definicion epica una historia",
                        "PO Definicion historia",
                        "PO Casos exito",
                        "PO Definicion mejora tecnica",
                        "PO Definicion spike"
                    ]
                )
                
                usar_todo_doc = st.checkbox(
                    "Usar documento completo",
                    value=False,
                    help="Si está marcado, usa todo el documento. Si no, solo fragmentos relevantes"
                )
            
            if st.button("✨ Generar Work Items", use_container_width=True):
                if not instruccion_generacion:
                    st.error("❌ Debes proporcionar instrucciones de generación")
                else:
                    # Decidir contexto
                    if usar_todo_doc:
                        contexto_doc = f"**Documento completo:**\n\n{st.session_state.doc_content[:10000]}"
                        if len(st.session_state.doc_content) > 10000:
                            contexto_doc += "\n\n[Documento truncado por longitud]"
                    else:
                        # Buscar fragmentos relevantes según instrucción
                        with st.spinner("🔍 Buscando fragmentos relevantes..."):
                            resultados = buscar_chunks_similares(
                                instruccion_generacion,
                                st.session_state.doc_chunks,
                                st.session_state.doc_embeddings,
                                st.session_state.embedding_model,
                                top_k=st.session_state.get('doc_top_k', 3)
                            )
                        contexto_doc = construir_contexto_documento(resultados)
                    
                    # Obtener template
                    template_content = get_template(template_generacion)
                    
                    # Preparar prompt combinado
                    prompt_final = f"""Contexto del documento:
{contexto_doc}

Instrucciones:
{instruccion_generacion}

Plantilla a seguir:
{template_content}

Genera el/los work item(s) solicitados siguiendo la plantilla proporcionada y basándote en el contenido del documento."""
                    
                    payload = {
                        "model": st.session_state.model,
                        "messages": [
                            {"role": "user", "content": prompt_final}
                        ]
                    }
                    
                    if st.session_state.include_temp:
                        payload["temperature"] = st.session_state.temperature
                    if st.session_state.include_tokens:
                        payload["max_tokens"] = st.session_state.max_tokens
                    
                    with st.spinner("🤖 Frida está generando los work items..."):
                        resultado_generacion = call_ia(payload)
                    
                    # Mostrar resultado
                    st.success("✅ Work items generados")
                    st.markdown("### Resultado:")
                    st.markdown(resultado_generacion)
                    
                    # Botón para copiar
                    st.text_area(
                        "Copiar resultado",
                        value=resultado_generacion,
                        height=300
                    )
